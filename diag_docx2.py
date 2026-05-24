"""
Full DOCX table analysis - understand the repeating structure of all 120 tables.
"""
import re
from pathlib import Path

try:
    import docx as _docx
except ImportError:
    print("python-docx not installed")
    exit()

docx_path = Path("data/April-26 TRC-9001/RUN NO. D(IGP-BSL) DN LINE/TrackParametersUrgentReport_2026_04_20_20_31_24_361.docx")
doc = _docx.Document(docx_path)

print(f"Total tables: {len(doc.tables)}")
print()

# Show first 12 tables to understand repeating block structure
for t_idx, tbl in enumerate(doc.tables[:12]):
    print(f"--- Table {t_idx} ---")
    for r_idx, row in enumerate(tbl.rows):
        cells = [c.text.replace('\xa0', ' ').strip() for c in row.cells]
        print(f"  Row {r_idx}: {cells}")
    print()

print("\n=== ANALYZING ALL TABLE PATTERNS ===")
# Count unique row-0 patterns
from collections import Counter
patterns = Counter()
for i, tbl in enumerate(doc.tables):
    if tbl.rows:
        key = str(len(tbl.rows)) + "rows_" + str(len(tbl.columns)) + "cols"
        patterns[key] += 1
        
for pattern, count in patterns.most_common():
    print(f"  {pattern}: {count} tables")
