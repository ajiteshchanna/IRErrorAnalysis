"""
docx_tqi_pipeline.py
====================
DOCX-Native Railway TQI / KM-wise Summary ETL Pipeline
=======================================================

Reads Word (.docx) files, extracts ONLY structured tables (ignoring all
paragraphs, images, headers, footers, inspection notes), validates them as
railway TQI/geometry summary tables, and stores each embedded Word table as
its own SQLite table in railway.db.

Naming convention:
    Word file   : 9001_RunC.docx
    Table 0     : tqi_9001_runc_t00
    Table 1     : tqi_9001_runc_t01

Architecture:
    process_folder()
        → extract_doc_tables()
            → is_valid_railway_table()
            → extract_metadata()
            → clean_columns()
            → clean_table()
            → standardize_columns()
        → store_to_sqlite()
        → register_processed_file()
    → run_analytics_views()

Usage:
    python docx_tqi_pipeline.py <folder_path> [db_path]

    Examples:
        python docx_tqi_pipeline.py ./data ./railway.db
        python docx_tqi_pipeline.py "data/April-26 TRC-9001" railway.db
"""

from __future__ import annotations

import hashlib
import logging
import re
import sqlite3
import sys
from datetime import datetime
from pathlib import Path
from typing import Iterator

import pandas as pd

# ── Optional import guard ────────────────────────────────────────────────────
try:
    import docx as _docx_lib
except ImportError:
    print(
        "[FATAL] python-docx is not installed.\n"
        "        Run:  pip install python-docx\n"
    )
    sys.exit(1)

# ── Shared column utilities ──────────────────────────────────────────────────
from column_utils import clean_column_name, EXTENDED_COLUMN_ALIASES

# ════════════════════════════════════════════════════════════════════════════
#  Constants & Configuration
# ════════════════════════════════════════════════════════════════════════════

DB_PATH        = "railway.db"
LOG_FORMAT     = "%(asctime)s [%(levelname)s] %(message)s"
TABLE_PREFIX   = "tqi_"           # prefix for all tables created by this pipeline
MAX_META_SCAN  = 15               # paragraphs to scan before/around a table for metadata
MIN_COLS       = 2                # minimum columns for a table to be considered
MIN_ROWS       = 1                # minimum data rows (after header) for a table to be kept
MAX_NULL_RATIO = 0.85             # columns with > 85% nulls are dropped
MAX_FILE_MB    = 50               # skip DOCX files larger than this (MB). Override with --max-file-mb N
                                  # Set to 0 to disable the limit entirely.

# Keywords that must appear in column headers for a table to be "railway valid"
# At least MIN_KEYWORD_HITS of these must be present (case-insensitive).
RAILWAY_KEYWORDS = {
    "km", "tqi", "uni", "ali", "spd", "speed", "s.no", "sno", "sr.no",
    "from", "to", "index", "gauge", "twist", "cross", "level", "section",
    "trc", "run", "alignment", "unevenness", "curvature"
}
MIN_KEYWORD_HITS = 2

# ─── Fallback (non-TQI) structured-table extraction ──────────────────────────
# Tables that fail the RAILWAY_KEYWORDS check are re-evaluated against these
# relaxed criteria. Passing tables receive the REPORT_TABLE_PREFIX instead of
# TABLE_PREFIX so they are easy to distinguish in the database.

REPORT_TABLE_PREFIX  = "rpt_"  # prefix for generic structured-table output
FALLBACK_MIN_COLS    = 2       # minimum column count
FALLBACK_MIN_ROWS    = 2       # minimum data-row count (after header)
FALLBACK_MIN_DENSITY = 0.20    # fraction of cells that must be non-empty

# Ordered list of (filename-keyword, short-report-type-tag) pairs.
# First match wins. Used for logging and possible future routing.
REPORT_TYPE_MAP: list[tuple[str, str]] = [
    ("urgentpeak",             "urgent_peak"),      # large image-heavy reports (usually size-skipped)
    ("wearurgent",             "wear_urgent"),
    ("trackparametersurgent",  "urgent_params"),
    ("trackparameterskm",      "kmwise_params"),
    ("trackparametersoffline", "offline_params"),
    ("optionsreport",          "options"),
    ("featuresnearuml",        "features_uml"),
    ("kmwisesummary",          "kmwise_summary"),
    ("trc",                    "trc_summary"),
    ("tqi",                    "tqi_summary"),
]

# Canonical column alias map — built from the shared EXTENDED_COLUMN_ALIASES.
# Kept as a module-level dict so existing code that references COLUMN_ALIASES
# continues to work without modification.
COLUMN_ALIASES: dict[str, str] = dict(EXTENDED_COLUMN_ALIASES)

# Metadata column names — appended to every row
# Includes the original 7 fields plus 7 new fields for richer context.
META_COLS = [
    # Original fields (preserved for backward compat with app.py queries)
    "trc_no",
    "run_no",
    "section",
    "direction",
    "run_date",
    # New clean geography & report-header fields
    "railway",
    "division",
    "route",
    "rt_code",
    "file_name",
    "start_km",
    "line",
    # System fields
    "source_file",
    "word_table_index",
]

# ════════════════════════════════════════════════════════════════════════════
#  Logging Setup
# ════════════════════════════════════════════════════════════════════════════

logging.basicConfig(
    level=logging.INFO,
    format=LOG_FORMAT,
    handlers=[
        logging.StreamHandler(sys.stdout),
    ],
)
log = logging.getLogger("docx_tqi")

# Force UTF-8 on Windows stdout to avoid cp1252 encoding errors
if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except AttributeError:
        pass  # Python < 3.7 or non-reconfigurable stream


# ════════════════════════════════════════════════════════════════════════════
#  1. File Utilities
# ════════════════════════════════════════════════════════════════════════════

def hash_file(path: Path, chunk: int = 1 << 20) -> str:
    """Compute SHA-256 of a file in streaming chunks. Returns hex digest."""
    h = hashlib.sha256()
    try:
        with open(path, "rb") as fh:
            while True:
                data = fh.read(chunk)
                if not data:
                    break
                h.update(data)
    except OSError as exc:
        log.warning(f"  [HASH] Cannot read '{path.name}': {exc}")
        return ""
    return h.hexdigest()


def generate_table_name(filepath: Path, table_index: int) -> str:
    """
    Convert a DOCX filepath + table index into a valid SQLite table name.

    Rules:
      - Strip extension, lowercase, replace non-alphanumerics with '_'
      - Collapse consecutive underscores, strip leading/trailing '_'
      - Prefix TABLE_PREFIX  (default: 'tqi_')
      - Suffix '_t<NN>'  where NN is zero-padded table index
      - If stem starts with a digit, insert 'f' between prefix and stem
      - Truncate stem portion to 50 characters (to keep total < 64)

    Example:
        "9001_RunC.docx", table 2  →  "tqi_f9001_runc_t02"
        "TRC Summary BD-IGP.docx", table 0  →  "tqi_trc_summary_bd_igp_t00"
    """
    stem = filepath.stem.strip()
    name = stem.lower()
    name = re.sub(r"[^a-z0-9]", "_", name)
    name = re.sub(r"_+", "_", name).strip("_") or "doc"

    if name[0].isdigit():
        name = "f" + name          # avoid leading digit in table name

    stem_part = name[:50]
    suffix    = f"_t{table_index:02d}"
    return f"{TABLE_PREFIX}{stem_part}{suffix}"


# Two-stage patterns used by _extract_trc_from_path.
#
# Stage A: TRC-labelled number — "TRC-9001", "TRC_9001", "TRC 9001", "TRC-NO-9001"
#   Allows an optional non-digit word (e.g. "NO") between TRC and the number.
#   Most reliable because the label is explicit.
#
# Stage B: bare 4+ digit number — used for folder names only (not filenames),
#   since filenames almost always contain date/timestamp digits (2026, 04, 20).
#   Uses (?<!\d) / (?!\d) to catch numbers after underscores (run_9001).
_PATH_TRC_LABELED_RE = re.compile(
    r"TRC[\-_\s]?(?:[A-Z]{1,6}[\-_\s])?(\d{4,})",
    re.IGNORECASE,
)
_PATH_TRC_BARE_RE = re.compile(r"(?<!\d)(\d{4,})(?!\d)")



def _extract_trc_from_path(doc_path: Path) -> str:
    """
    Extract a TRC number from the file path using this priority order:

      Stage A — TRC-labelled number in filename or any parent folder
          TRC_9001_RunD.docx          -> 9001  (filename, labelled)
          /data/April-26 TRC-9001/f  -> 9001  (folder, labelled)
          /data/TRC-9001/file.docx   -> 9001  (folder, labelled)

      Stage B — Bare 4+ digit number in a parent folder name only
          /data/9001/report.docx     -> 9001  (folder, bare)
          /path/run_9001/file.docx   -> 9001  (folder, bare)
          (filenames excluded to avoid year/date false positives like '2026')

      Stage C — Returns "" (caller falls back to DOCX header metadata)

    Returns:
        String of digits (e.g. '9001'), or "" if nothing reliable is found.
    """
    # ── Stage A: TRC-labelled match in filename or any ancestor folder ────────
    m = _PATH_TRC_LABELED_RE.search(doc_path.stem)
    if m:
        log.debug("  [TRC-PATH] trc_no=%r from filename (labelled) %r", m.group(1), doc_path.name)
        return m.group(1)

    for part in reversed(doc_path.parts[:-1]):
        m = _PATH_TRC_LABELED_RE.search(part)
        if m:
            log.debug("  [TRC-PATH] trc_no=%r from folder (labelled) %r", m.group(1), part)
            return m.group(1)

    # ── Stage B: bare 4+ digit number — folder names only ────────────────────
    # Filenames are intentionally excluded here because they commonly contain
    # date/time digits (2026, 0420, 2031) that are indistinguishable from TRC
    # numbers by length alone.
    for part in reversed(doc_path.parts[:-1]):
        m = _PATH_TRC_BARE_RE.search(part)
        if m:
            log.debug("  [TRC-PATH] trc_no=%r from folder (bare) %r", m.group(1), part)
            return m.group(1)

    # ── Stage C: not found ────────────────────────────────────────────────────
    return ""



def _detect_report_type(filepath: Path) -> str:
    """
    Classify a DOCX file by its filename using REPORT_TYPE_MAP.
    Returns a short tag string (e.g. 'wear_urgent', 'options') or 'report'.
    """
    name_lower = filepath.stem.lower()
    for keyword, rtype in REPORT_TYPE_MAP:
        if keyword in name_lower:
            return rtype
    return "report"


def generate_rpt_table_name(filepath: Path, table_index: int) -> str:
    """
    Generate a REPORT_TABLE_PREFIX (rpt_) table name for non-TQI structured tables.

    Follows the same rules as generate_table_name() but uses REPORT_TABLE_PREFIX
    so the resulting tables are clearly distinct from tqi_* tables.

    Example:
        "WearUrgentReport_2026_04_19.docx", table 0  →  "rpt_wearurgentreport_2026_04_19_t00"
        "OptionsReport_2026_04_19.docx",    table 2  →  "rpt_optionsreport_2026_04_19_t02"
    """
    stem = filepath.stem.strip()
    name = stem.lower()
    name = re.sub(r"[^a-z0-9]", "_", name)
    name = re.sub(r"_+", "_", name).strip("_") or "doc"

    if name[0].isdigit():
        name = "f" + name

    stem_part = name[:50]
    suffix    = f"_t{table_index:02d}"
    return f"{REPORT_TABLE_PREFIX}{stem_part}{suffix}"



# ════════════════════════════════════════════════════════════════════════════
#  2. DOCX Table Extraction
# ════════════════════════════════════════════════════════════════════════════

def _cell_text(cell) -> str:
    """Return clean text from a python-docx Table cell."""
    try:
        return cell.text.replace("\xa0", " ").replace("\x00", "").strip()
    except Exception:
        return ""


def _table_to_raw_rows(tbl) -> list[list[str]]:
    """
    Convert a python-docx Table object into a list of string rows.
    Handles merged cells by repeating the text (forward-fill across merges).
    """
    rows: list[list[str]] = []
    for row in tbl.rows:
        cells = [_cell_text(cell) for cell in row.cells]
        rows.append(cells)
    return rows


def _find_header_row(rows: list[list[str]]) -> int:
    """
    Heuristic: the header row is the first row where ≥ MIN_KEYWORD_HITS
    of RAILWAY_KEYWORDS appear in the row content.

    Returns the index of the header row, or -1 if not found.
    """
    for i, row in enumerate(rows):
        row_text = " ".join(row).lower()
        hits = sum(1 for kw in RAILWAY_KEYWORDS if kw in row_text)
        if hits >= MIN_KEYWORD_HITS:
            return i
    return -1


def _rows_to_dataframe(rows: list[list[str]], header_idx: int) -> pd.DataFrame | None:
    """
    Build a DataFrame from raw string rows using the identified header row.
    Rows above the header (title rows) are discarded.
    The header row becomes column names; subsequent rows become data.
    """
    header_row = rows[header_idx]
    data_rows  = rows[header_idx + 1:]

    if not data_rows:
        return None

    # Normalise width — pad or trim each data row to match header width
    n_cols = len(header_row)
    normed = []
    for row in data_rows:
        if len(row) < n_cols:
            row = row + [""] * (n_cols - len(row))
        normed.append(row[:n_cols])

    # Replace empty column names with positional fallbacks
    cols = []
    for i, h in enumerate(header_row):
        h = h.strip()
        cols.append(h if h else f"col_{i}")

    df = pd.DataFrame(normed, columns=cols)
    return df


def extract_doc_tables(doc_path: Path) -> list[tuple[int, pd.DataFrame]]:
    """
    Open a .docx file and extract every table as a DataFrame.

    Primary path   — tables whose rows contain >= MIN_KEYWORD_HITS railway keywords.
                     The keyword-bearing row becomes the column header.

    Fallback path  — tables where no railway-keyword header is found, but the table
                     has >= FALLBACK_MIN_ROWS+1 rows and >= FALLBACK_MIN_COLS columns.
                     Row 0 is used as the column header unconditionally.

    Returns:
        List of (table_index, DataFrame) for tables that could be parsed.
        Table index is the 0-based position of the table in the document.

    Raises nothing — all errors are caught and logged.
    """
    results: list[tuple[int, pd.DataFrame]] = []

    try:
        doc = _docx_lib.Document(str(doc_path))
    except Exception as exc:
        log.error(f"  [OPEN] Cannot open '{doc_path.name}': {exc}")
        return results

    n_tables = len(doc.tables)
    log.info(f"  Found {n_tables} table(s) in document.")

    for idx, tbl in enumerate(doc.tables):
        try:
            raw_rows = _table_to_raw_rows(tbl)

            if len(raw_rows) < 2:
                log.debug(f"  Table {idx}: too few rows ({len(raw_rows)}) — skip.")
                continue

            header_idx = _find_header_row(raw_rows)

            if header_idx != -1:
                # Primary path: railway keyword header found
                df = _rows_to_dataframe(raw_rows, header_idx)
                if df is None or df.empty:
                    log.debug(f"  Table {idx}: empty after primary header extraction — skip.")
                    continue
                results.append((idx, df))
                log.info(
                    f"  Table {idx}: extracted {len(df)} raw row(s), "
                    f"{len(df.columns)} col(s)."
                )
            else:
                # Fallback path: no railway keyword header — use row 0 as header
                n_cols = len(raw_rows[0]) if raw_rows else 0
                n_data_rows = len(raw_rows) - 1
                if n_cols < FALLBACK_MIN_COLS or n_data_rows < FALLBACK_MIN_ROWS:
                    log.debug(
                        f"  Table {idx}: no railway header, too small for fallback "
                        f"({n_data_rows} data rows x {n_cols} cols) — skip."
                    )
                    continue
                df = _rows_to_dataframe(raw_rows, 0)
                if df is None or df.empty:
                    log.debug(f"  Table {idx}: empty after fallback extraction — skip.")
                    continue
                results.append((idx, df))
                log.info(
                    f"  Table {idx}: fallback-extracted {len(df)} row(s), "
                    f"{len(df.columns)} col(s) (no railway keywords in headers)."
                )

        except Exception as exc:
            log.error(f"  Table {idx}: extraction error: {exc}", exc_info=False)

    return results



# ════════════════════════════════════════════════════════════════════════════
#  3. Table Validation
# ════════════════════════════════════════════════════════════════════════════

def is_valid_railway_table(df: pd.DataFrame) -> bool:
    """
    Validate that a DataFrame looks like a railway TQI/geometry table.

    Checks:
      1. At least MIN_COLS columns.
      2. At least MIN_ROWS data rows.
      3. At least MIN_KEYWORD_HITS railway keywords present in column names.
      4. Not all values are empty strings / NaN.
    """
    if len(df.columns) < MIN_COLS:
        return False
    if len(df) < MIN_ROWS:
        return False

    col_text = " ".join(str(c).lower() for c in df.columns)
    hits = sum(1 for kw in RAILWAY_KEYWORDS if kw in col_text)
    if hits < MIN_KEYWORD_HITS:
        return False

    # Check that the table has at least some non-empty data
    non_empty = df.apply(
        lambda col: col.map(lambda v: str(v).strip() not in ("", "nan", "None"))
    ).values.any()
    return bool(non_empty)


def is_valid_fallback_table(df: pd.DataFrame) -> bool:
    """
    Relaxed validation for non-TQI structured Word tables.

    A table passes if:
      1. It has at least FALLBACK_MIN_COLS columns.
      2. It has at least FALLBACK_MIN_ROWS data rows.
      3. At least FALLBACK_MIN_DENSITY fraction of cells are non-empty.

    Used as a second-chance validator when is_valid_railway_table() returns False.
    Passing tables are stored with the REPORT_TABLE_PREFIX (rpt_) naming scheme.
    """
    if len(df.columns) < FALLBACK_MIN_COLS:
        return False
    if len(df) < FALLBACK_MIN_ROWS:
        return False

    total_cells = df.shape[0] * df.shape[1]
    if total_cells == 0:
        return False

    non_empty_count = df.apply(
        lambda col: col.map(
            lambda v: str(v).strip() not in ("", "nan", "None", "none")
        )
    ).values.sum()

    density = non_empty_count / total_cells
    return density >= FALLBACK_MIN_DENSITY


# ════════════════════════════════════════════════════════════════════════════
#  4. Metadata Extraction
# ════════════════════════════════════════════════════════════════════════════

# A lookahead that stops any value-capture at the next metadata keyword or
# cell boundary.  Used by railway / division / section patterns to prevent
# "value" from consuming neighbouring key:value tokens when cells are joined
# into a single string with spaces.
_META_STOP = r"(?=\s*(?:\||RAILWAY|DIVISION|SECTION|START\s*KM|LINE|ROUTE|RT[-\s]?CODE|FILE\s*NAME|RUN\s*NO|DATE|TRC\s*NO|$))"

# ── Multi-pattern lists for fields with inconsistent DOCX formats ─────────
# trc_no and run_date appear in many format variants across Indian Railways
# DOCX reports.  Each field is stored as a list of compiled patterns tried
# in priority order by _scan_text_for_metadata() — first match wins.
#
# TRC NO variants seen in the wild:
#   TRC No: TRC-9001  |  TRC NO - 9001  |  TRC No:9001  |  TRC NO: TRC9001
#   TRC NO:- 9001     |  TRC NUMBER 9001
#
# Date label variants seen in the wild:
#   Date of Rec: 18.04.2026  |  RUN Date: 20/04/2026
#   date: 13.04.2026         |  Run Date : 2026-04-20
_TRC_PATTERNS: list = [
    # P1 — Labelled (NO/NUMBER) + optional 'TRC-' prefix in value, digits only
    #   Matches: TRC NO: 9001 | TRC NO - 9001 | TRC No: TRC-9001 | TRC NO:TRC9001
    re.compile(
        r"TRC\s*(?:NO\.?|NUMBER)\s*[:\-\.\=\s\u2011\u2013\u2014]*\s*"
        r"(?:TRC[\-_\s]?)?(\d{3,})",
        re.IGNORECASE,
    ),
    # P2 — Labelled (NO/NUMBER) + alphanumeric value, extracts trailing digits
    #   Matches: TRC NO: TRC9001 | TRC NO: RT10201321 (extracts last 4+ digits)
    re.compile(
        r"TRC\s*(?:NO\.?|NUMBER)\s*[:\-\.\=\s\u2011\u2013\u2014]*\s*"
        r"[A-Z]*[\-_]?(\d{3,})",
        re.IGNORECASE,
    ),
    # P3 — Bare 'TRC-9001' or 'TRC 9001' without a NO label
    #   Matches: (TRC-9001) | TRC 9001 — used only when P1/P2 find nothing
    re.compile(
        r"TRC[\-_\s](\d{4,})",
        re.IGNORECASE,
    ),
]

_DATE_PATTERNS: list = [
    # P1 — 'Date of Rec' (most specific label, highest priority)
    #   Matches: Date of Rec: 18.04.2026 | Date of Rec - 18/04/2026
    re.compile(
        r"Date\s*of\s*Rec\s*[:\-]?\s*(\d{1,4}[/\-\.]\d{1,2}[/\-\.]\d{1,4})",
        re.IGNORECASE,
    ),
    # P2 — 'RUN Date' / 'Run Date' (two-word label)
    #   Matches: RUN Date: 20/04/2026 | Run Date : 2026-04-20
    re.compile(
        r"RUN\s+Date\s*[:\-]?\s*(\d{1,4}[/\-\.]\d{1,2}[/\-\.]\d{1,4})",
        re.IGNORECASE,
    ),
    # P3 — Bare 'Date:' or 'DATE:' (single-word label, word-boundary guarded)
    #   Matches: date: 13.04.2026 | DATE: 20.04.2026
    #   (?<!\w) prevents matching 'RunDate:' mid-word
    re.compile(
        r"(?<!\w)Date\s*[:\-]\s*(\d{1,4}[/\-\.]\d{1,2}[/\-\.]\d{1,4})",
        re.IGNORECASE,
    ),
    # P4 — 'Date' with optional separator (fallback — no mandatory colon)
    #   Matches: Date 13.04.2026 (space separator only)
    re.compile(
        r"(?<!\w)Date\s+(\d{1,4}[/\-\.]\d{1,2}[/\-\.]\d{1,4})",
        re.IGNORECASE,
    ),
]

_META_PATTERNS = {
    # trc_no and run_date now point to their pattern lists (handled specially
    # by _scan_text_for_metadata — first match in the list wins).
    "trc_no":    _TRC_PATTERNS,
    # RUN NO — single letter or alphanumeric (e.g. D, A, B1)
    "run_no":    re.compile(
        r"Run\s*(?:No\.?|Number)\s*[:\-\.\s]?\s*([A-Z0-9][A-Z0-9\-/()]*)",
        re.IGNORECASE
    ),
    # SECTION — stop at the next metadata keyword or cell boundary
    "section":   re.compile(
        r"Section\s*[:\-]?\s*([A-Z0-9][A-Z0-9\-/\s]*?)" + _META_STOP,
        re.IGNORECASE
    ),
    "direction": re.compile(
        r"Direction\s*[:\-]?\s*(UP|DN|DOWN|UP\s*LINE|DN\s*LINE|DOWN\s*LINE)",
        re.IGNORECASE
    ),
    # run_date now points to its pattern list.
    "run_date":  _DATE_PATTERNS,
    # ── Extended geography / report-header fields ──────────────────────────
    # All stop at the next metadata keyword to prevent cross-field bleed.
    # railway / division also accept 'Railways:' / 'Division:' label variants
    # (some OptionsReport headers use that casing/spelling)
    "railway":   re.compile(
        # Lazy *? ensures we stop at the next metadata keyword (_META_STOP lookahead)
        r"RAILWAYS?\s*[:\-]?\s*([A-Za-z][A-Za-z\s]*?)" + _META_STOP,
        re.IGNORECASE
    ),
    "division":  re.compile(
        r"DIVISIONS?\s*[:\-]?\s*([A-Za-z0-9][A-Za-z0-9\-/\s]*?)" + _META_STOP,
        re.IGNORECASE
    ),
    "route":     re.compile(
        r"ROUTE\s*[:\-]?\s*([^\s|\n]+)",
        re.IGNORECASE
    ),
    "rt_code":   re.compile(
        r"RT[\-\s]?CODE\s*[:\-]?\s*([^\s|\n]+)",
        re.IGNORECASE
    ),
    "file_name": re.compile(
        r"FILE\s*NAME\s*[:\-]?\s*([^\s|\n]+)",
        re.IGNORECASE
    ),
    "start_km":  re.compile(
        r"START\s*KM\s*[:\-]?\s*([\d.]+)",
        re.IGNORECASE
    ),
    "line":      re.compile(
        r"\bLINE\s*[:\-]?\s*(UP|DN|DOWN|IInd|IIIrd|IVth|3rd|4th)\b",
        re.IGNORECASE
    ),
}

# Final-resort date fallback — pure numeric date anywhere in text, no label
# required.  Only fires if all _DATE_PATTERNS fail.  Both DD.MM.YYYY and
# YYYY.MM.DD forms are accepted; 4-digit year is mandatory to avoid false hits.
_DATE_FALLBACK = re.compile(
    r"\b(\d{1,4}[/\-\.]\d{1,2}[/\-\.]\d{1,4})\b"
)


def _normalise_date(raw: str) -> str:
    """
    Normalise a raw date string to ISO YYYY-MM-DD format.

    Accepts separator styles:  .  /  -
    Accepts formats:
        DD[sep]MM[sep]YYYY   →   e.g. 20.04.2026  →  2026-04-20
        YYYY[sep]MM[sep]DD   →   e.g. 2026-04-20  →  2026-04-20 (unchanged)
        DD[sep]MM[sep]YY     →   e.g. 20.04.26    →  2026-04-20 (2000s assumed)

    Returns the original raw string unchanged if parsing fails, so no data
    is ever lost even for unexpected formats.
    """
    if not raw:
        return raw
    # Normalise all separators to hyphen for uniform splitting
    normalised = re.sub(r"[./]", "-", raw.strip())
    parts = normalised.split("-")
    if len(parts) != 3:
        return raw
    a, b, c = parts
    try:
        if len(c) == 4:
            # DD-MM-YYYY
            day, month, year = int(a), int(b), int(c)
        elif len(a) == 4:
            # YYYY-MM-DD
            year, month, day = int(a), int(b), int(c)
        elif len(c) == 2:
            # DD-MM-YY  →  assume 2000s
            day, month, year = int(a), int(b), 2000 + int(c)
        else:
            return raw  # unrecognised layout — keep as-is
        # Sanity check
        if not (1 <= month <= 12 and 1 <= day <= 31 and 1900 <= year <= 2100):
            return raw
        return f"{year:04d}-{month:02d}-{day:02d}"
    except (ValueError, TypeError):
        return raw


def _scan_text_for_metadata(text: str) -> dict:
    """
    Scan a single block of text for all metadata fields.
    Returns a dict with found values (or "" for missing fields).

    Fields whose value in _META_PATTERNS is a LIST of patterns (trc_no,
    run_date) are tried in priority order — first match wins.  All other
    fields use a single compiled pattern as before.

    run_date is always normalised to YYYY-MM-DD before returning.
    """
    result: dict[str, str] = {}
    for key, pattern_or_list in _META_PATTERNS.items():
        if isinstance(pattern_or_list, list):
            # Multi-pattern field: try each pattern in order, use first match
            matched = ""
            for pat in pattern_or_list:
                m = pat.search(text)
                if m:
                    matched = m.group(1).strip()
                    break
            result[key] = matched
        else:
            m = pattern_or_list.search(text)
            result[key] = m.group(1).strip() if m else ""

    # Final-resort date fallback — pure date anywhere, no label required.
    # Accepts both DD.MM.YYYY and YYYY.MM.DD; 4-digit year prevents false hits.
    if not result.get("run_date"):
        m = _DATE_FALLBACK.search(text)
        if m:
            result["run_date"] = m.group(1)

    # Always normalise run_date to ISO YYYY-MM-DD
    if result.get("run_date"):
        result["run_date"] = _normalise_date(result["run_date"])

    return result


# Label→field mapping for cell-pair scanning.
# Key  = normalised label text (lowercase, stripped)
# Value = metadata field name
# Used when a table row has ["TRC NO", "9001"] in adjacent cells instead of
# "TRC NO: 9001" in a single cell.
_CELL_LABEL_MAP: dict[str, str] = {
    "trc no":       "trc_no",
    "trc no.":      "trc_no",
    "trc number":   "trc_no",
    "date":         "run_date",
    "run date":     "run_date",
    "date of rec":  "run_date",
    "run no":       "run_no",
    "run number":   "run_no",
    "run no.":      "run_no",
    "section":      "section",
    "sections":     "section",
    "railway":      "railway",
    "railways":     "railway",
    "division":     "division",
    "route":        "route",
    "rt-code":      "rt_code",
    "rt code":      "rt_code",
    "file name":    "file_name",
    "start km":     "start_km",
    "line":         "line",
    "direction":    "direction",
}


def _extract_cell_pairs(row_cells: list[str], meta: dict) -> None:
    """
    Scan consecutive cell pairs (and triples) in a table row and fill
    metadata gaps.  Handles two common DOCX header table layouts:

    Two-cell layout (value immediately after label):
        Cell 0         Cell 1
        ───────────    ──────────
        "TRC NO"       "9001"
        "DATE"         "20.04.2026"

    Three-cell layout (colon separator in its own cell):
        Cell 0         Cell 1    Cell 2
        ───────────    ───────   ──────────
        "TRC No."      ":"       "9001"
        "Railway"      ":"       "Central"

    Modifies `meta` in-place; only fills gaps (does not overwrite existing).
    """
    n = len(row_cells)
    for i in range(n - 1):
        label_raw  = row_cells[i].strip()
        label      = label_raw.rstrip(":.\u2013\u2014\u2011-").strip().lower()
        value_cell = row_cells[i + 1].strip()

        # Three-cell triple: if next cell is just a separator, look one further
        if value_cell in (":", "-", "=", "|") and i + 2 < n:
            value_cell = row_cells[i + 2].strip()

        if not label or not value_cell:
            continue
        field = _CELL_LABEL_MAP.get(label)
        if field and not meta.get(field):
            if field == "run_date":
                normalised = _normalise_date(value_cell)
                if re.match(r"\d{4}-\d{2}-\d{2}", normalised):
                    meta[field] = normalised
            elif field == "trc_no":
                # Value may be '9001' or 'TRC-9001' or just digits
                m = re.search(r"(\d+)", value_cell)
                if m:
                    meta[field] = m.group(1)
            else:
                meta[field] = value_cell


def extract_metadata(doc, table_index: int) -> dict:
    """
    Extract metadata directly from the raw DOCX Document object.

    Extracts 12 distinct fields:
      Original : trc_no, run_no, section, direction, run_date
      New      : railway, division, route, rt_code, file_name, start_km, line

    Strategy (in priority order — earlier fills win):
      1. Scan all document paragraphs.  Metadata in paragraph text before the
         first table is captured here (e.g. title rows, report headers).
      2. Scan ALL tables (not just the target table):
           a. Joined row text  → regex-based extraction.
           b. Adjacent cell-pair scan → handles label/value split across cells.
      3. Scan the parent folder name for leftover section/direction/run hints.

    NOTE:  Strategy 2b is the key fix for TRC NO / DATE NULL values.
    Indian Railways DOCX header tables often store metadata as two-cell rows:
        ["TRC NO", "9001"] or ["DATE", "20.04.2026"]
    python-docx joins these into a single string without the colon separator,
    so the standard regex fails.  The cell-pair scan catches this layout.

    Args:
        doc         : python-docx Document object (opened BEFORE table extraction)
        table_index : 0-based index of the table being processed (currently
                      unused in extraction but retained for API compatibility)

    Returns:
        dict with 12 keys, one per metadata field.
    """
    meta: dict[str, str] = {
        "trc_no":    "",
        "run_no":    "",
        "section":   "",
        "direction": "",
        "run_date":  "",
        # Extended fields
        "railway":   "",
        "division":  "",
        "route":     "",
        "rt_code":   "",
        "file_name": "",
        "start_km":  "",
        "line":      "",
    }

    # ── Strategy 1: build ONE combined full-document text blob ───────────────
    #
    # Combine ALL paragraphs AND ALL table cell text into a single searchable
    # string.  This ensures metadata buried anywhere in the document — whether
    # in a top-level paragraph, a header table cell, or a merged cell — is
    # captured by a single regex pass.
    #
    # doc.paragraphs = top-level only (NOT inside table cells).
    # doc.tables[*].rows[*].cells[*].text = cell paragraphs joined by \n.
    # Both sources are needed because DOCX stores them separately.
    all_text_parts: list[str] = []
    try:
        for para in doc.paragraphs:
            t = para.text.replace("\xa0", " ").replace("\x00", "").strip()
            if t:
                all_text_parts.append(t)
    except Exception as exc:
        log.debug(f"  [META] Paragraph scan error: {exc}")

    try:
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    t = cell.text.replace("\xa0", " ").replace("\x00", "").strip()
                    if t:
                        all_text_parts.append(t)
    except Exception as exc:
        log.debug(f"  [META] Table cell text scan error: {exc}")

    if all_text_parts:
        full_doc_text = " | ".join(all_text_parts)
        blob_meta = _scan_text_for_metadata(full_doc_text)
        for k, v in blob_meta.items():
            if v:
                meta[k] = v
        log.debug(
            "  [META-BLOB] trc_no=%r run_date=%r (from full-doc text blob)",
            meta.get("trc_no"), meta.get("run_date"),
        )

    # ── Strategy 2: per-row cell-pair scan (catches split label|value cells) ──
    #
    # The regex blob above joins cells with ' | ' which usually lets regex
    # match.  But some reports have the label in one cell and the bare value
    # in the adjacent cell with NO textual separator at all:
    #   Cell0="TRC NO"  Cell1="9001"
    # When the cells are joined as "TRC NO | 9001" the regex still works, but
    # we keep the cell-pair scan as an additional safety net for edge cases
    # and to fill any remaining gaps.
    try:
        for tbl in doc.tables:
            scan_rows = list(tbl.rows)[:MAX_META_SCAN]
            for row in scan_rows:
                cells = [_cell_text(c) for c in row.cells]

                # 2a: regex on the full joined row text (gap-fill only)
                row_text = " ".join(cells)
                row_meta = _scan_text_for_metadata(row_text)
                for k, v in row_meta.items():
                    if v and not meta.get(k):
                        meta[k] = v

                # 2b: cell-pair scan — catches ["TRC NO", "9001"] split cells
                _extract_cell_pairs(cells, meta)

                # Early exit once all critical fields are filled
                if (
                    meta["trc_no"] and meta["run_date"]
                    and meta["run_no"] and meta["section"]
                ):
                    break
    except Exception as exc:
        log.debug(f"  [META] Table header scan error: {exc}")

    # ── Strategy 3: folder-name fallback ─────────────────────────────────────
    try:
        folder_name = ""
        if hasattr(doc, "_path"):          # python-docx <1.x private attr
            folder_name = Path(doc._path).parent.name
        if folder_name:
            folder_meta = _scan_text_for_metadata(folder_name)
            for k, v in folder_meta.items():
                if v and not meta.get(k):
                    meta[k] = v
    except Exception:
        pass

    log.debug(
        "  [META] trc_no=%r run_date=%r run_no=%r section=%r "
        "railway=%r division=%r",
        meta["trc_no"], meta["run_date"], meta["run_no"], meta["section"],
        meta["railway"], meta["division"],
    )
    return meta


# ════════════════════════════════════════════════════════════════════════════
#  5. Column Cleaning & Standardisation
# ════════════════════════════════════════════════════════════════════════════

def _raw_to_snake(raw: str) -> str:
    """
    Convert a raw column header string to a safe snake_case identifier.

    Steps:
      1. Strip whitespace
      2. Lowercase
      3. Replace all non-alphanumeric characters with '_'
      4. Collapse consecutive underscores
      5. Strip leading/trailing underscores
    """
    s = raw.strip().lower()
    s = re.sub(r"[^a-z0-9]", "_", s)
    s = re.sub(r"_+", "_", s).strip("_")
    return s or "col"


def standardize_columns(df: pd.DataFrame) -> pd.DataFrame:
    """
    Apply the canonical COLUMN_ALIASES map to column names.

    First converts each column to snake_case via _raw_to_snake(), then
    looks up the result (and the original lowercased name) in COLUMN_ALIASES.
    Falls back to the snake_case version if no alias is found.
    """
    new_cols: list[str] = []
    for col in df.columns:
        raw_lower = col.strip().lower()
        snake     = _raw_to_snake(col)

        # Try original lowercase first, then snake_case version
        alias = COLUMN_ALIASES.get(raw_lower) or COLUMN_ALIASES.get(snake) or snake
        new_cols.append(alias)

    df.columns = new_cols
    return df


def clean_columns(df: pd.DataFrame) -> pd.DataFrame:
    """
    Clean and deduplicate DataFrame column names.

    Steps:
      1. Apply standardize_columns() (alias map + snake_case)
      2. Drop columns named 'col', 'col_N', or that are empty strings
      3. Deduplicate column names by appending _2, _3, ...
    """
    df = standardize_columns(df)

    # Drop positional placeholder columns (e.g., col_0, col_5)
    keep = [c for c in df.columns if not re.fullmatch(r"col(_\d+)?", c)]
    df = df[keep] if keep else df

    # Drop fully-empty column name
    df = df.loc[:, df.columns != ""]

    # Deduplicate
    seen: dict[str, int] = {}
    new_cols: list[str] = []
    for col in df.columns:
        if col in seen:
            seen[col] += 1
            new_cols.append(f"{col}_{seen[col]}")
        else:
            seen[col] = 0
            new_cols.append(col)
    df.columns = new_cols
    return df


def clean_table(df: pd.DataFrame) -> pd.DataFrame:
    """
    Remove noise rows and columns from an extracted table DataFrame.

    Steps:
      1. Drop fully-empty rows and columns.
      2. Drop columns with > MAX_NULL_RATIO fraction of empty/NaN values.
      3. Drop rows where the FIRST column is not numeric (catches embedded
         sub-headers or total/summary rows within the data body).
      4. Reset index.
    """
    # Step 1: drop all-empty rows and columns
    df = df.replace({"": None})
    df = df.dropna(how="all").reset_index(drop=True)
    df = df.dropna(how="all", axis=1)

    if df.empty:
        return df

    # Step 2: drop mostly-null columns
    null_ratio = df.isnull().mean()
    df = df.loc[:, null_ratio <= MAX_NULL_RATIO]

    if df.empty:
        return df

    # Step 3: filter data rows — first column should be numeric (S.No or KM)
    def _is_numeric(val) -> bool:
        try:
            float(str(val).strip().replace(",", ""))
            return True
        except (ValueError, TypeError):
            return False

    first_col = df.columns[0]
    mask = df[first_col].map(_is_numeric)

    # Allow non-numeric first column if < 30% of rows are numeric
    # (some tables start with KM range strings like "663+000")
    km_range_pat = re.compile(r"^\d+[\+\-\.]\d+$")
    def _is_km_range(val) -> bool:
        return bool(km_range_pat.match(str(val).strip()))

    km_mask = df[first_col].map(_is_km_range)
    valid_mask = mask | km_mask
    data_df = df[valid_mask].reset_index(drop=True)

    # If filtering removed everything, return original (don't over-filter)
    if data_df.empty:
        return df.reset_index(drop=True)

    return data_df


# ════════════════════════════════════════════════════════════════════════════
#  6. SQLite Storage
# ════════════════════════════════════════════════════════════════════════════

def _ensure_processed_files_table(conn: sqlite3.Connection) -> None:
    """
    Create the processed_files registry (with file_hash column) if it
    doesn't exist. Safely adds file_hash column to existing tables.
    """
    conn.execute("""
        CREATE TABLE IF NOT EXISTS processed_files (
            id               INTEGER PRIMARY KEY AUTOINCREMENT,
            source_filename  TEXT    NOT NULL,
            generated_table  TEXT    NOT NULL,
            file_type        TEXT,
            processed_date   TEXT,
            last_modified    TEXT,
            record_count     INTEGER DEFAULT 0,
            status           TEXT    DEFAULT 'ok',
            file_hash        TEXT,
            UNIQUE(source_filename, generated_table)
        )
    """)
    # Ensure file_hash column exists on older registry tables
    existing_cols = [
        row[1] for row in conn.execute("PRAGMA table_info(processed_files)")
    ]
    if "file_hash" not in existing_cols:
        try:
            conn.execute(
                "ALTER TABLE processed_files ADD COLUMN file_hash TEXT"
            )
            log.info("  [REGISTRY] Added 'file_hash' column to processed_files.")
        except sqlite3.OperationalError:
            pass  # Already exists in a concurrent run

    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_pf_filename "
        "ON processed_files(source_filename)"
    )
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_pf_table "
        "ON processed_files(generated_table)"
    )
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_pf_hash "
        "ON processed_files(file_hash)"
    )
    conn.commit()


def _is_already_processed(db_path: str, file_hash: str) -> bool:
    """
    Return True if a file with this SHA-256 hash has already been processed
    successfully (status = 'ok') in this pipeline (file_type = 'docx_tqi').
    """
    if not file_hash:
        return False
    try:
        with sqlite3.connect(db_path) as conn:
            _ensure_processed_files_table(conn)
            row = conn.execute(
                "SELECT id FROM processed_files "
                "WHERE file_hash = ? AND status = 'ok' AND file_type = 'docx_tqi'",
                (file_hash,)
            ).fetchone()
            return row is not None
    except Exception:
        return False


def _get_existing_columns(conn: sqlite3.Connection, table_name: str) -> list[str]:
    """Return list of column names for an existing table, or [] if table absent."""
    row = conn.execute(
        "SELECT name FROM sqlite_master WHERE type='table' AND name=?",
        (table_name,)
    ).fetchone()
    if row is None:
        return []
    return [r[1] for r in conn.execute(f'PRAGMA table_info("{table_name}")')]


def store_to_sqlite(
    df: pd.DataFrame,
    table_name: str,
    db_path: str = DB_PATH,
) -> int:
    """
    Write a DataFrame to SQLite.

    - First write  → creates the table via pandas to_sql.
    - Later writes → only inserts columns that already exist in the table
                     (schema-stable append).
    - Sanitizes all string values (removes null bytes).
    - Returns the number of rows written (0 on error).
    """
    if df.empty:
        log.warning(f"  [STORE] '{table_name}' has no rows — skipping write.")
        return 0

    # Sanitize: remove null bytes from string columns
    for col in df.select_dtypes(include="object").columns:
        df[col] = df[col].map(
            lambda v: str(v).replace("\x00", "").strip()
            if pd.notna(v) else v
        )

    try:
        with sqlite3.connect(db_path) as conn:
            existing = _get_existing_columns(conn, table_name)

            if not existing:
                # First write — create the table
                df.to_sql(table_name, conn, if_exists="append", index=False)
                log.info(
                    f"  [STORE] Created '{table_name}' and inserted "
                    f"{len(df)} row(s) × {len(df.columns)} col(s)."
                )
            else:
                # Subsequent write — only insert common columns
                common = [c for c in df.columns if c in existing]
                if not common:
                    log.warning(
                        f"  [STORE] No common columns for '{table_name}' "
                        f"(df has {list(df.columns)}, table has {existing}) — skip."
                    )
                    return 0
                df[common].to_sql(table_name, conn, if_exists="append", index=False)
                log.info(
                    f"  [STORE] Appended {len(df)} row(s) -> '{table_name}' "
                    f"({len(common)} cols matched)."
                )
        return len(df)

    except Exception as exc:
        log.error(f"  [STORE] DB write error for '{table_name}': {exc}", exc_info=True)
        return 0


def register_processed_file(
    db_path: str,
    source_filename: str,
    generated_table: str,
    last_modified: str,
    record_count: int,
    file_hash: str,
    status: str = "ok",
) -> None:
    """
    Upsert one row into the processed_files registry.
    Uses file_type = 'docx_tqi' to distinguish from the old pipeline.
    """
    processed_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    try:
        with sqlite3.connect(db_path) as conn:
            _ensure_processed_files_table(conn)
            conn.execute(
                """
                INSERT INTO processed_files
                    (source_filename, generated_table, file_type,
                     processed_date, last_modified, record_count, status, file_hash)
                VALUES (?, ?, 'docx_tqi', ?, ?, ?, ?, ?)
                ON CONFLICT(source_filename, generated_table) DO UPDATE SET
                    processed_date = excluded.processed_date,
                    last_modified  = excluded.last_modified,
                    record_count   = record_count + excluded.record_count,
                    status         = excluded.status,
                    file_hash      = excluded.file_hash
                """,
                (
                    source_filename, generated_table,
                    processed_date, last_modified,
                    record_count, status, file_hash,
                ),
            )
            conn.commit()
        log.info(
            "  [REGISTRY] '%s' -> '%s' (%d rows, status=%s)",
            source_filename, generated_table, record_count, status,
        )
    except Exception as exc:
        log.error(f"  [REGISTRY] Failed to register '{source_filename}': {exc}")


# ════════════════════════════════════════════════════════════════════════════
#  7. Analytics Views (Compatibility Layer)
# ════════════════════════════════════════════════════════════════════════════

def run_analytics_views(db_path: str = DB_PATH) -> None:
    """
    Create or refresh SQL views for analytics compatibility.

    TQI views (existing, unchanged):
      v_tqi_all      — UNION ALL of every tqi_* table
      v_worst_tqi_s  — bottom 100 rows by tqi_s  (worst short-chord)
      v_worst_tqi_l  — bottom 100 rows by tqi_l  (worst long-chord)
      v_worst_tqi_c  — bottom 100 rows by tqi_c  (worst composite)
      v_tqi_by_run   — aggregate stats per (section, direction, run_no, run_date)
      v_tqi_by_km    — aggregate stats per km_from (cross-run comparison)

    Report views (new — rpt_* tables from fallback extraction):
      v_rpt_all      — UNION ALL of every rpt_* table with common columns
      v_rpt_by_file  — row counts and metadata grouped by source_file
    """
    log.info("[VIEWS] Building analytics views ...")

    try:
        with sqlite3.connect(db_path) as conn:

            # ═══════════════════════════════════════════════════════════════
            #  TQI VIEWS (existing logic — unchanged)
            # ═══════════════════════════════════════════════════════════════

            tqi_tables = [
                r[0]
                for r in conn.execute(
                    "SELECT name FROM sqlite_master "
                    "WHERE type='table' AND name LIKE 'tqi_%' "
                    "AND name NOT LIKE 'tqi_view_%'"
                )
            ]

            if not tqi_tables:
                log.info("  [VIEWS] No tqi_* tables found yet — skipping TQI views.")
            else:
                log.info(f"  [VIEWS] Found {len(tqi_tables)} tqi_* table(s).")

                all_cols_set: dict[str, None] = {}
                for tbl in tqi_tables:
                    for r in conn.execute(f'PRAGMA table_info("{tbl}")'):
                        all_cols_set[r[1]] = None
                all_cols = list(all_cols_set.keys())

                def _select_expr(tbl: str) -> str:
                    tbl_cols = {r[1] for r in conn.execute(f'PRAGMA table_info("{tbl}")')}
                    parts = [
                        f'"{c}"' if c in tbl_cols else f'NULL AS "{c}"'
                        for c in all_cols
                    ]
                    return (
                        f'SELECT {", ".join(parts)}, \'{tbl}\' AS _tqi_source '
                        f'FROM "{tbl}"'
                    )

                union_sql = "\nUNION ALL\n".join(_select_expr(t) for t in tqi_tables)
                conn.execute("DROP VIEW IF EXISTS v_tqi_all")
                conn.execute(f"CREATE VIEW v_tqi_all AS\n{union_sql}")
                log.info("  [VIEWS] Created v_tqi_all")

                tqi_s_exists = "tqi_s" in all_cols

                if tqi_s_exists:
                    conn.execute("DROP VIEW IF EXISTS v_worst_tqi_s")
                    conn.execute("""
                        CREATE VIEW v_worst_tqi_s AS
                        SELECT * FROM v_tqi_all
                        WHERE tqi_s IS NOT NULL AND tqi_s != ''
                          AND CAST(tqi_s AS REAL) IS NOT NULL
                        ORDER BY CAST(tqi_s AS REAL) ASC
                        LIMIT 100
                    """)
                    log.info("  [VIEWS] Created v_worst_tqi_s")

                if "tqi_l" in all_cols:
                    conn.execute("DROP VIEW IF EXISTS v_worst_tqi_l")
                    conn.execute("""
                        CREATE VIEW v_worst_tqi_l AS
                        SELECT * FROM v_tqi_all
                        WHERE tqi_l IS NOT NULL AND tqi_l != ''
                        ORDER BY CAST(tqi_l AS REAL) ASC
                        LIMIT 100
                    """)
                    log.info("  [VIEWS] Created v_worst_tqi_l")

                if "tqi_c" in all_cols:
                    conn.execute("DROP VIEW IF EXISTS v_worst_tqi_c")
                    conn.execute("""
                        CREATE VIEW v_worst_tqi_c AS
                        SELECT * FROM v_tqi_all
                        WHERE tqi_c IS NOT NULL AND tqi_c != ''
                        ORDER BY CAST(tqi_c AS REAL) ASC
                        LIMIT 100
                    """)
                    log.info("  [VIEWS] Created v_worst_tqi_c")

                agg_cols = []
                for col in ["tqi_s", "tqi_l", "tqi_c", "uni_1", "ali_1", "spd"]:
                    if col in all_cols:
                        agg_cols.append(
                            f"  AVG(CAST({col} AS REAL)) AS avg_{col}, "
                            f"MIN(CAST({col} AS REAL)) AS min_{col}, "
                            f"MAX(CAST({col} AS REAL)) AS max_{col}"
                        )

                if agg_cols and "section" in all_cols:
                    conn.execute("DROP VIEW IF EXISTS v_tqi_by_run")
                    conn.execute(f"""
                        CREATE VIEW v_tqi_by_run AS
                        SELECT
                            section, direction, trc_no, run_no, run_date,
                            COUNT(*) AS km_count,
                            {', '.join(agg_cols)}
                        FROM v_tqi_all
                        GROUP BY section, direction, trc_no, run_no, run_date
                        ORDER BY run_date DESC
                    """)
                    log.info("  [VIEWS] Created v_tqi_by_run")

                if tqi_s_exists and "km_from" in all_cols:
                    conn.execute("DROP VIEW IF EXISTS v_tqi_by_km")
                    conn.execute("""
                        CREATE VIEW v_tqi_by_km AS
                        SELECT
                            km_from, km_to, section, direction,
                            COUNT(DISTINCT run_no) AS run_count,
                            AVG(CAST(tqi_s AS REAL)) AS avg_tqi_s,
                            MIN(CAST(tqi_s AS REAL)) AS min_tqi_s,
                            MAX(CAST(tqi_s AS REAL)) AS max_tqi_s
                        FROM v_tqi_all
                        WHERE km_from IS NOT NULL AND km_from != ''
                        GROUP BY km_from, km_to, section, direction
                        ORDER BY avg_tqi_s ASC
                    """)
                    log.info("  [VIEWS] Created v_tqi_by_km")

            # ═══════════════════════════════════════════════════════════════
            #  RPT VIEWS (new — for rpt_* fallback-extracted tables)
            # ═══════════════════════════════════════════════════════════════

            rpt_tables = [
                r[0]
                for r in conn.execute(
                    "SELECT name FROM sqlite_master "
                    "WHERE type='table' AND name LIKE 'rpt_%'"
                )
            ]

            if not rpt_tables:
                log.info("  [VIEWS] No rpt_* tables found yet — skipping RPT views.")
            else:
                log.info(f"  [VIEWS] Found {len(rpt_tables)} rpt_* table(s).")

                # Common columns guaranteed to exist on every rpt_* table
                # (because _process_single_docx always attaches them)
                RPT_COMMON_COLS = [
                    "source_file", "word_table_index", "report_type",
                    "trc_no", "run_no", "section", "direction", "run_date",
                ]

                def _rpt_select_expr(tbl: str) -> str:
                    tbl_cols = {r[1] for r in conn.execute(f'PRAGMA table_info("{tbl}")')}
                    parts = [
                        f'"{c}"' if c in tbl_cols else f'NULL AS "{c}"'
                        for c in RPT_COMMON_COLS
                    ]
                    return (
                        f'SELECT {", ".join(parts)}, \'{tbl}\' AS _rpt_source '
                        f'FROM "{tbl}"'
                    )

                rpt_union = "\nUNION ALL\n".join(
                    _rpt_select_expr(t) for t in rpt_tables
                )
                conn.execute("DROP VIEW IF EXISTS v_rpt_all")
                conn.execute(f"CREATE VIEW v_rpt_all AS\n{rpt_union}")
                log.info("  [VIEWS] Created v_rpt_all")

                # v_rpt_by_file — quick summary per source document
                conn.execute("DROP VIEW IF EXISTS v_rpt_by_file")
                conn.execute("""
                    CREATE VIEW v_rpt_by_file AS
                    SELECT
                        source_file,
                        report_type,
                        section,
                        direction,
                        trc_no,
                        run_no,
                        run_date,
                        COUNT(*) AS row_count,
                        COUNT(DISTINCT _rpt_source) AS table_count
                    FROM v_rpt_all
                    GROUP BY source_file, report_type, section, direction,
                             trc_no, run_no, run_date
                    ORDER BY source_file
                """)
                log.info("  [VIEWS] Created v_rpt_by_file")

            # ═══════════════════════════════════════════════════════════════
            #  COMPATIBILITY VIEWS (new — exception-report & meta summary)
            # ═══════════════════════════════════════════════════════════════

            # v_exception_report_all — UNION ALL of any table that carries the
            # exception-report signature columns (parameter + recorded_value + loc_km)
            try:
                exc_tables = [
                    r[0]
                    for r in conn.execute(
                        "SELECT name FROM sqlite_master WHERE type='table' "
                        "AND name NOT IN ('processed_files')"
                    )
                    if all(
                        col in {r2[1] for r2 in conn.execute(f'PRAGMA table_info("{r[0]}")')}
                        for col in ("parameter", "recorded_value", "loc_km")
                    )
                ]
                if exc_tables:
                    exc_common = [
                        "parameter", "threshold", "recorded_value",
                        "loc_km", "loc_meter",
                        "trc_no", "date", "run_no", "route", "rt_code",
                        "railway", "division", "section", "start_km", "line",
                        "source_file",
                    ]

                    def _exc_sel(tbl: str) -> str:
                        tbl_cols = {r[1] for r in conn.execute(f'PRAGMA table_info("{tbl}")')}
                        parts = [
                            f'"{c}"' if c in tbl_cols else f'NULL AS "{c}"'
                            for c in exc_common
                        ]
                        return f'SELECT {", ".join(parts)}, \'{tbl}\' AS _source FROM "{tbl}"'

                    exc_union = "\nUNION ALL\n".join(_exc_sel(t) for t in exc_tables)
                    conn.execute("DROP VIEW IF EXISTS v_exception_report_all")
                    conn.execute(f"CREATE VIEW v_exception_report_all AS\n{exc_union}")
                    log.info(f"  [VIEWS] Created v_exception_report_all ({len(exc_tables)} table(s))")
                else:
                    log.info("  [VIEWS] No exception-report tables found — skipping v_exception_report_all.")
            except Exception as exc_ex:
                log.warning(f"  [VIEWS] v_exception_report_all build failed: {exc_ex}")

            # v_meta_summary — one row per (source_file, trc_no, run_no) from
            # both tqi_* and rpt_* tables that carry the extended metadata columns
            try:
                meta_tables = [
                    r[0]
                    for r in conn.execute(
                        "SELECT name FROM sqlite_master WHERE type='table' "
                        "AND name NOT IN ('processed_files')"
                    )
                    if {r2[1] for r2 in conn.execute(f'PRAGMA table_info("{r[0]}")')}
                    .issuperset({"trc_no", "source_file"})
                ]
                if meta_tables:
                    meta_common = [
                        "trc_no", "run_no", "section", "direction", "run_date",
                        "railway", "division", "route", "rt_code",
                        "start_km", "line", "source_file",
                    ]

                    def _meta_sel(tbl: str) -> str:
                        tbl_cols = {r[1] for r in conn.execute(f'PRAGMA table_info("{tbl}")')}
                        parts = [
                            f'"{c}"' if c in tbl_cols else f'NULL AS "{c}"'
                            for c in meta_common
                        ]
                        return f'SELECT {", ".join(parts)}, \'{tbl}\' AS _meta_source FROM "{tbl}"'

                    meta_union = "\nUNION ALL\n".join(_meta_sel(t) for t in meta_tables)
                    conn.execute("DROP VIEW IF EXISTS v_meta_summary")
                    conn.execute(f"""
                        CREATE VIEW v_meta_summary AS
                        SELECT
                            trc_no, run_no, section, direction, run_date,
                            railway, division, route, rt_code, start_km, line,
                            source_file,
                            COUNT(*) AS row_count,
                            COUNT(DISTINCT _meta_source) AS table_count
                        FROM ({meta_union})
                        GROUP BY trc_no, run_no, section, direction, run_date,
                                 railway, division, route, rt_code, start_km, line,
                                 source_file
                        ORDER BY run_date DESC, source_file
                    """)
                    log.info(f"  [VIEWS] Created v_meta_summary ({len(meta_tables)} table(s))")
                else:
                    log.info("  [VIEWS] No tables with trc_no+source_file found — skipping v_meta_summary.")
            except Exception as meta_ex:
                log.warning(f"  [VIEWS] v_meta_summary build failed: {meta_ex}")

            conn.commit()
            log.info("[VIEWS] All analytics views created successfully.")

    except Exception as exc:
        log.error(f"[VIEWS] Failed to build analytics views: {exc}", exc_info=True)


# ════════════════════════════════════════════════════════════════════════════
#  8. Main Pipeline
# ════════════════════════════════════════════════════════════════════════════

def _discover_docx_files(
    folder: Path,
    max_file_mb: float = MAX_FILE_MB,
) -> Iterator[Path]:
    """
    Yield all .docx files recursively under folder.

    Skips:
      - Temp/lock files (names starting with '~')
      - Files larger than max_file_mb MB (when max_file_mb > 0)
        These are typically UrgentPeakWithImage reports with thousands of
        embedded photo tables that are not KM-wise TQI summaries.
    """
    limit_bytes = int(max_file_mb * 1024 * 1024) if max_file_mb > 0 else 0
    for f in sorted(folder.rglob("*.docx")):
        if f.name.startswith("~"):
            log.warning(f"  [SKIP] Temp/lock file: {f.name}")
            continue
        if limit_bytes > 0:
            size_mb = f.stat().st_size / (1024 * 1024)
            if f.stat().st_size > limit_bytes:
                log.warning(
                    "  [SKIP] File too large (%.1f MB > %.0f MB limit): %s",
                    size_mb, max_file_mb, f.name,
                )
                continue
        yield f


def _process_single_docx(
    doc_path: Path,
    db_path: str,
    summary: dict,
    force: bool = False,
) -> tuple[int, int]:
    """
    Full pipeline for one DOCX file.

    Returns (tables_written, total_rows_inserted).
    """
    log.info(f"\n{'=' * 62}")
    log.info(f"[FILE] {doc_path.name}")
    log.info(f"  PATH: {doc_path}")

    # ── Hash check ───────────────────────────────────────────────────────────
    file_hash = hash_file(doc_path)
    last_mod  = str(doc_path.stat().st_mtime)

    if not force and _is_already_processed(db_path, file_hash):
        log.info("  [SKIP] File hash already processed — use --force to reprocess.")
        return 0, 0

    # ── Open document ────────────────────────────────────────────────────────
    try:
        doc = _docx_lib.Document(str(doc_path))
    except Exception as exc:
        log.error(f"  [OPEN] Cannot open document: {exc}")
        register_processed_file(
            db_path, doc_path.name, f"{TABLE_PREFIX}error",
            last_mod, 0, file_hash, "error"
        )
        return 0, 0

    # ── Extract metadata from the document ───────────────────────────────────
    log.info("  Extracting document metadata ...")

    # Resolve TRC number from the file path FIRST (most reliable source).
    # This is used to override the DOCX-extracted trc_no after extract_metadata()
    # runs, ensuring filename/folder takes priority over embedded header text.
    path_trc_no = _extract_trc_from_path(doc_path)
    if path_trc_no:
        log.info(f"  [TRC-PATH] trc_no={path_trc_no!r} (from filename/folder)")
    else:
        log.info("  [TRC-PATH] No TRC number found in filename/folder — will use DOCX metadata.")

    # ── Extract tables ────────────────────────────────────────────────────────
    log.info("  Extracting tables from document ...")
    raw_tables = extract_doc_tables(doc_path)

    if not raw_tables:
        log.warning(f"  [SKIP] No extractable tables found in '{doc_path.name}'.")
        register_processed_file(
            db_path, doc_path.name, f"{TABLE_PREFIX}skipped",
            last_mod, 0, file_hash, "skipped"
        )
        return 0, 0

    tables_written = 0
    total_rows     = 0

    for tbl_idx, raw_df in raw_tables:
        log.info(f"  .. Processing Word Table {tbl_idx} ..")

        # ── Primary path: TQI / railway-keyword validated table ───────────────
        if is_valid_railway_table(raw_df):
            log.info(f"    [TQI-VALID] Table {tbl_idx}: {len(raw_df)} rows × {len(raw_df.columns)} cols")

            # Extract metadata from DOCX (paragraphs + header tables)
            meta = extract_metadata(doc, tbl_idx)

            # ── TRC number override: filename/folder wins over DOCX text ─────
            # Priority 1/2 (path) beats Priority 3 (DOCX header).
            # If the path gave nothing, keep whatever DOCX extraction found.
            if path_trc_no:
                meta["trc_no"] = path_trc_no

            log.info(
                f"    Metadata: section='{meta['section']}' "
                f"trc='{meta['trc_no']}' run='{meta['run_no']}' "
                f"dir='{meta['direction']}' date='{meta['run_date']}'"
            )

            df = clean_columns(raw_df.copy())
            df = clean_table(df)
            if df.empty:
                log.warning(f"    [SKIP] Table {tbl_idx}: empty after cleaning.")
                continue

            log.info(f"    After cleaning: {len(df)} rows × {len(df.columns)} cols")
            log.info(f"    Columns: {list(df.columns)}")

            df["trc_no"]           = meta.get("trc_no",    "")
            df["run_no"]           = meta.get("run_no",    "")
            df["section"]          = meta.get("section",   "")
            df["direction"]        = meta.get("direction", "")
            df["run_date"]         = meta.get("run_date",  "")
            # New extended metadata fields
            df["railway"]          = meta.get("railway",   "")
            df["division"]         = meta.get("division",  "")
            df["route"]            = meta.get("route",     "")
            df["rt_code"]          = meta.get("rt_code",   "")
            df["file_name"]        = meta.get("file_name", "")
            df["start_km"]         = meta.get("start_km",  "")
            df["line"]             = meta.get("line",      "")
            df["source_file"]      = doc_path.name
            df["word_table_index"] = tbl_idx

            data_cols = [c for c in df.columns if c not in META_COLS]
            df = df[data_cols + [c for c in META_COLS if c in df.columns]]

            tbl_name = generate_table_name(doc_path, tbl_idx)
            log.info(f"    -> TQI SQLite table: '{tbl_name}'")

            rows_written = store_to_sqlite(df, tbl_name, db_path)
            if rows_written > 0:
                tables_written += 1
                total_rows     += rows_written
                summary[tbl_name] = summary.get(tbl_name, 0) + rows_written
                register_processed_file(
                    db_path, doc_path.name, tbl_name,
                    last_mod, rows_written, file_hash, "ok"
                )
            else:
                log.warning(f"    [WARN] Table {tbl_idx}: 0 rows written (TQI path).")

        # ── Fallback path: any structured table with ≥ MIN_COLS, MIN_ROWS, MIN_DENSITY ──
        elif is_valid_fallback_table(raw_df):
            report_type = _detect_report_type(doc_path)
            log.info(
                f"    [RPT-VALID] Table {tbl_idx} (type={report_type}): "
                f"{len(raw_df)} rows × {len(raw_df.columns)} cols — using fallback extraction."
            )

            # Extract metadata (same function — gracefully returns empty strings if not found)
            meta = extract_metadata(doc, tbl_idx)

            # ── TRC number override: filename/folder wins over DOCX text ─────
            if path_trc_no:
                meta["trc_no"] = path_trc_no

            # Use column cleaning but skip the COLUMN_ALIASES alias step (raw columns preserved)
            df = clean_table(raw_df.copy())
            # Normalise column names to safe snake_case identifiers
            df = standardize_columns(df)
            df = clean_columns(df)

            if df.empty:
                log.warning(f"    [SKIP] Table {tbl_idx}: empty after fallback cleaning.")
                continue

            log.info(f"    After cleaning: {len(df)} rows × {len(df.columns)} cols")
            log.info(f"    Columns: {list(df.columns)}")

            # Attach same metadata columns as TQI path for dashboard compatibility
            df["trc_no"]           = meta.get("trc_no",    "")
            df["run_no"]           = meta.get("run_no",    "")
            df["section"]          = meta.get("section",   "")
            df["direction"]        = meta.get("direction", "")
            df["run_date"]         = meta.get("run_date",  "")
            # New extended metadata fields
            df["railway"]          = meta.get("railway",   "")
            df["division"]         = meta.get("division",  "")
            df["route"]            = meta.get("route",     "")
            df["rt_code"]          = meta.get("rt_code",   "")
            df["file_name"]        = meta.get("file_name", "")
            df["start_km"]         = meta.get("start_km",  "")
            df["line"]             = meta.get("line",      "")
            df["source_file"]      = doc_path.name
            df["word_table_index"] = tbl_idx
            df["report_type"]      = report_type  # extra: classification tag

            # RPT_META_COLS = META_COLS + report_type
            rpt_meta = META_COLS + ["report_type"]
            data_cols = [c for c in df.columns if c not in rpt_meta]
            all_meta  = [c for c in rpt_meta if c in df.columns]
            df = df[data_cols + all_meta]

            tbl_name = generate_rpt_table_name(doc_path, tbl_idx)
            log.info(f"    -> RPT SQLite table: '{tbl_name}'")

            rows_written = store_to_sqlite(df, tbl_name, db_path)
            if rows_written > 0:
                tables_written += 1
                total_rows     += rows_written
                summary[tbl_name] = summary.get(tbl_name, 0) + rows_written
                register_processed_file(
                    db_path, doc_path.name, tbl_name,
                    last_mod, rows_written, file_hash, "ok"
                )
            else:
                log.warning(f"    [WARN] Table {tbl_idx}: 0 rows written (RPT path).")

        else:
            log.info(
                f"    [SKIP] Table {tbl_idx}: failed both TQI and fallback validation "
                f"(cols: {list(raw_df.columns)[:5]}..., rows: {len(raw_df)})."
            )

    if tables_written == 0:
        log.warning(
            f"  [RESULT] '{doc_path.name}': no valid tables extracted (TQI or RPT)."
        )
        register_processed_file(
            db_path, doc_path.name, f"{TABLE_PREFIX}no_valid_tables",
            last_mod, 0, file_hash, "skipped"
        )

    return tables_written, total_rows


def process_folder(
    folder_path: str,
    db_path: str = DB_PATH,
    force: bool = False,
    max_file_mb: float = MAX_FILE_MB,
) -> None:
    """
    Main entry point. Recursively processes all .docx files in folder_path.

    Args:
        folder_path : Path to the folder containing .docx files
        db_path     : Path to railway.db (created if absent)
        force       : If True, reprocess files even if hash matches
        max_file_mb : Skip DOCX files larger than this size in MB (0 = no limit).
                      Defaults to MAX_FILE_MB (50 MB). Use 0 to process ALL files
                      including very large UrgentPeakWithImage reports.
    """
    folder = Path(folder_path)

    if not folder.exists():
        log.error(f"[FATAL] Folder does not exist: {folder_path}")
        return
    if not folder.is_dir():
        log.error(f"[FATAL] Not a directory: {folder_path}")
        return

    # ── Initialise DB ────────────────────────────────────────────────────────
    with sqlite3.connect(db_path) as conn:
        conn.execute("PRAGMA journal_mode=WAL")
        _ensure_processed_files_table(conn)
    log.info(f"[DB] Using database: {db_path}")

    # ── Discover files ───────────────────────────────────────────────────────
    docx_files = list(_discover_docx_files(folder, max_file_mb=max_file_mb))

    log.info("=" * 62)
    log.info("[SCAN] Folder    : %s", folder_path)
    log.info("[SCAN] Found     : %d .docx file(s)", len(docx_files))
    log.info("[SCAN] Max size  : %.0f MB (0 = unlimited)", max_file_mb)
    if force:
        log.info("[SCAN] --force flag active: reprocessing all files")
    log.info("=" * 62)

    if not docx_files:
        log.warning("[DONE] No .docx files found. Nothing to process.")
        return

    # ── Process each file ────────────────────────────────────────────────────
    summary: dict[str, int] = {}
    total_files_ok      = 0
    total_files_skipped = 0
    total_tables        = 0
    total_rows          = 0

    for doc_path in docx_files:
        try:
            tables_ok, rows_ok = _process_single_docx(
                doc_path, db_path, summary, force=force
            )
            if tables_ok > 0:
                total_files_ok += 1
                total_tables   += tables_ok
                total_rows     += rows_ok
            else:
                total_files_skipped += 1
        except Exception as exc:
            log.error(
                f"[ERROR] Unexpected error processing '{doc_path.name}': {exc}",
                exc_info=True,
            )
            total_files_skipped += 1

    # ── Build analytics views ────────────────────────────────────────────────
    if total_tables > 0:
        run_analytics_views(db_path)

    # ── Final Summary ────────────────────────────────────────────────────────
    log.info(f"\n{'=' * 62}")
    log.info("DOCX TQI PIPELINE — COMPLETE")
    log.info(f"  Files scanned    : {len(docx_files)}")
    log.info(f"  Files processed  : {total_files_ok}")
    log.info(f"  Files skipped    : {total_files_skipped}")
    log.info(f"  Tables extracted : {total_tables}")
    log.info(f"  Rows inserted    : {total_rows}")
    if summary:
        log.info("  Tables written:")
        for tbl, rows in sorted(summary.items()):
            log.info(f"    {tbl:<52} {rows:>6} rows")
    log.info(f"  Database         : {db_path}")
    log.info("=" * 62)


# ════════════════════════════════════════════════════════════════════════════
#  CLI Entry Point
# ════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    args = sys.argv[1:]

    if not args or args[0] in ("-h", "--help"):
        print(__doc__)
        print("Usage:   python docx_tqi_pipeline.py <folder_path> [db_path] [--force] [--max-file-mb N]")
        print("")
        print("Options:")
        print("  --force            Reprocess files even if already in the registry")
        print("  --max-file-mb N    Skip DOCX files larger than N MB (default: 50)")
        print("                     Use --max-file-mb 0 to process ALL files (slow for large reports)")
        print("")
        print("Examples:")
        print("  python docx_tqi_pipeline.py ./data railway.db")
        print("  python docx_tqi_pipeline.py ./data railway.db --force")
        print("  python docx_tqi_pipeline.py ./data railway.db --max-file-mb 100")
        print("  python docx_tqi_pipeline.py ./data railway.db --max-file-mb 0  # no limit")
        sys.exit(0)

    # Parse positional and optional args
    positional = [a for a in args if not a.startswith("--")]
    _folder    = positional[0] if len(positional) > 0 else "."
    _db        = positional[1] if len(positional) > 1 else DB_PATH
    _force     = "--force" in args

    # Parse --max-file-mb N
    _max_mb = MAX_FILE_MB
    if "--max-file-mb" in args:
        idx = args.index("--max-file-mb")
        if idx + 1 < len(args):
            try:
                _max_mb = float(args[idx + 1])
            except ValueError:
                print(f"[ERROR] --max-file-mb expects a number, got: {args[idx + 1]}")
                sys.exit(1)

    process_folder(_folder, _db, force=_force, max_file_mb=_max_mb)
