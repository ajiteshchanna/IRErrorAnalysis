"""
docx_to_xlsx_converter.py
=========================
Converts Indian Railways TRC DOCX files into XLSX files that are compatible
with the existing railway_pipeline.py EXCEPTION REPORT block-detection logic.

Three DOCX formats are supported:

  Format A — KM-wise Job Card  (TrackParametersKmJobCard_*.docx)
    • Contains N tables (one per KM), each with:
        row 3 : Railway | Division | Section | Line | KM
        row 5-6: header — Block | SD | KM | Metre | UML Param | UML Val | CBML Param | CBML Val | Track Feature
        row 7+ : data rows
    • Converted to a flat XLSX with one sheet per file.
    • Uses lxml for direct XML parsing (avoids docx.table crash on missing <w:tblGrid>).

  Format B — Urgent / Offline / KmReport  (TrackParametersUrgentReport_*.docx,
             TrackParametersOfflineReport_*.docx, TrackParametersKmReport_*.docx)
    • Repeating 4-table-per-block structure (30 blocks × 4 = 120 tables):
        Table 4k+0 : Metadata row 1 — TRC NO, DATE, RUN NO, ROUTE, RT-CODE, FILE NAME
        Table 4k+1 : Metadata row 2 — RAILWAY, DIVISION, SECTION, START KM, LINE
        Table 4k+2 : 2-row column header — PARAMETER / THRESHOLD / RECORDED VALUE / LOC→KM / LOC→METER
        Table 4k+3 : Data rows         — actual exception peak records
    • Each block may have different START KM (the section changes per block).
    • Metadata is extracted per block (START KM changes), then flattened into row columns.
    • section_spd is extracted from document paragraph text.

  Format C — Detail Ballast / Fastening / Rail Defect Report (all other *.docx)
    • Paragraph-based structure — unchanged from previous version.

Both Format A and B outputs prepend an "EXCEPTION REPORT" header row so that
detect_blocks() in railway_pipeline.py picks up the data automatically.

Usage (standalone):
    from docx_to_xlsx_converter import convert_docx_to_xlsx, is_conversion_current
    xlsx_path = convert_docx_to_xlsx(Path("report.docx"), Path("data/"))
"""

import re
import logging
from pathlib import Path

import pandas as pd

log = logging.getLogger(__name__)

# ── Namespace shorthand ───────────────────────────────────────────────────────
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ── Metadata field names ──────────────────────────────────────────────────────
_META_FIELDS = [
    "trc_no", "run_date", "run_no", "route", "rt_code", "file_name",
    "railway", "division", "section", "start_km", "line", "section_spd",
]

# ── Compiled regex for each metadata field ────────────────────────────────────
# Handles both colon and hyphen separators, flexible whitespace.
_META_RE = {
    "trc_no":      re.compile(r"TRC\s*NO\s*[:\-]\s*([A-Z0-9]+)", re.IGNORECASE),
    "run_date":    re.compile(
        r"DATE\s*[:\-]?\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{2,4})", re.IGNORECASE),
    "run_no":      re.compile(r"RUN\s*NO\.?\s*[:\-]?\s*([A-Z0-9]+)", re.IGNORECASE),
    "route":       re.compile(r"ROUTE\s*[:\-]?\s*([^\s|,\n]+)", re.IGNORECASE),
    "rt_code":     re.compile(r"RT[\-\s]?CODE\s*[:\-]?\s*([^\s|,\n]+)", re.IGNORECASE),
    "file_name":   re.compile(r"FILE\s*NAME\s*[:\-]?\s*([^\s|,\n]+)", re.IGNORECASE),
    "railway":     re.compile(r"RAILWAY\s*[:\-]?\s*([^|,\n]+?)(?:\s*\||\s*,|\s*$)", re.IGNORECASE),
    "division":    re.compile(r"DIVISION\s*[:\-]?\s*([^|,\n]+?)(?:\s*\||\s*,|\s*$)", re.IGNORECASE),
    "section":     re.compile(r"SECTION\s*[:\-]?\s*([^|,\n]+?)(?:\s*\||\s*,|\s*$)", re.IGNORECASE),
    "start_km":    re.compile(r"START\s*KM\s*[:\-]?\s*(\d[\d.]*)", re.IGNORECASE),
    "line":        re.compile(
        r"\bLINE\s*[:\-]?\s*(UP|DN|DOWN|IInd|IIIrd|IVth|3rd|4th)\b", re.IGNORECASE),
    "section_spd": re.compile(
        r"SECTION\s*SPD\s*(?:\(KMPH\))?\s*[:\-]?\s*(\d+)", re.IGNORECASE),
}

# Dash-only strings → treated as NULL
_DASH_RE = re.compile(r"^\s*[\-–—]+\s*$")


def _clean_meta_value(raw: str) -> str | None:
    """
    Sanitise a raw metadata string value.
    Returns None for empty, whitespace-only, or dash-only values.
    Strips trailing pipe chars, whitespace, and label-bleed artefacts.
    """
    if raw is None:
        return None
    v = str(raw).strip().strip("| \t")
    if not v or _DASH_RE.match(v):
        return None
    return v


# ═════════════════════════════════════════════════════════════════════════════
#  Helpers
# ═════════════════════════════════════════════════════════════════════════════

def is_conversion_current(docx_path: Path, xlsx_path: Path) -> bool:
    """Return True if xlsx_path exists and is at least as new as docx_path."""
    if not xlsx_path.exists():
        return False
    return xlsx_path.stat().st_mtime >= docx_path.stat().st_mtime


def _cell_text(tc_element) -> str:
    """Extract all text from a <w:tc> element, stripping non-breaking spaces."""
    texts = tc_element.findall(f".//{{{_W}}}t")
    return "".join(t.text or "" for t in texts).replace("\xa0", " ").strip()


def _table_rows(tbl_element) -> list[list[str]]:
    """Return list-of-lists of cell text for every row in a <w:tbl> element."""
    rows = []
    for tr in tbl_element.findall(f"{{{_W}}}tr"):
        cells = [_cell_text(tc) for tc in tr.findall(f"{{{_W}}}tc")]
        rows.append(cells)
    return rows


def _scan_text_for_meta(text: str) -> dict:
    """
    Scan a single text string for all metadata fields using _META_RE.
    Returns dict field→value (None for missing/dash-only).
    """
    result: dict = {}
    for field, pat in _META_RE.items():
        m = pat.search(text)
        raw = m.group(1).strip() if m else None
        result[field] = _clean_meta_value(raw)
    return result


def _merge_meta(base: dict, update: dict) -> dict:
    """Merge two meta dicts, only filling in gaps (base values are kept)."""
    out = dict(base)
    for k, v in update.items():
        if not out.get(k) and v:
            out[k] = v
    return out


def _normalise_date(raw: str | None) -> str | None:
    """
    Convert common date formats to YYYY-MM-DD.
    Handles: DD.MM.YYYY, DD/MM/YYYY, DD-MM-YYYY.
    Returns None if unparseable.
    """
    if not raw:
        return None
    m = re.match(r"(\d{1,2})[./\-](\d{1,2})[./\-](\d{2,4})", raw.strip())
    if not m:
        return raw  # already in some other format, return as-is
    d, mo, y = m.group(1), m.group(2), m.group(3)
    if len(y) == 2:
        y = "20" + y
    try:
        return f"{int(y):04d}-{int(mo):02d}-{int(d):02d}"
    except ValueError:
        return raw


def _write_xlsx(df: pd.DataFrame, out_path: Path, report_type: str,
                meta: dict, sheet_name: str = "Sheet1") -> Path:
    """
    Write df to XLSX in the EXCEPTION REPORT block format expected by
    the existing railway_pipeline.py extract_table() / detect_blocks() logic.

    Row layout (all written with header=False so pipeline reads header=None):
        [0] EXCEPTION REPORT — <report_type>  (detect_blocks trigger)
        [1] TRC No.: ...  RUN Date: ...  RUN No.: ...  ROUTE: ...  RT-CODE: ...  FILE NAME: ...
        [2] RAILWAY: ...  DIVISION: ...  SECTION: ...  START KM: ...  LINE: ...  SECTION SPD: ...
        [3] (blank separator)
        [4] col_0  col_1  col_2  ...   ← text header row
        [5] 1      val    val    ...   ← first data row
        ...
    """
    out_path.parent.mkdir(parents=True, exist_ok=True)

    def _v(key):
        return meta.get(key) or ""

    cols   = list(df.columns)
    n_cols = len(cols)
    empty  = [""] * (n_cols - 1)

    header_rows = [
        [f"EXCEPTION REPORT — {report_type}"] + empty,
        [f"TRC No.: {_v('trc_no')}  RUN Date: {_v('run_date')}  "
         f"RUN No.: {_v('run_no')}  ROUTE: {_v('route')}  "
         f"RT-CODE: {_v('rt_code')}  FILE NAME: {_v('file_name')}"] + empty,
        [f"RAILWAY: {_v('railway')}  DIVISION: {_v('division')}  "
         f"SECTION: {_v('section')}  START KM: {_v('start_km')}  "
         f"LINE: {_v('line')}  SECTION SPD: {_v('section_spd')}"] + empty,
        [""] * n_cols,   # blank separator
        cols,            # column-name row — picked up as table header
    ]

    header_df = pd.DataFrame(header_rows, columns=cols)
    combined  = pd.concat([header_df, df], ignore_index=True)

    with pd.ExcelWriter(out_path, engine="openpyxl") as writer:
        combined.to_excel(writer, sheet_name=sheet_name, index=False, header=False)

    log.info(f"    Wrote {len(df)} data rows → {out_path.name}")
    return out_path


# ═════════════════════════════════════════════════════════════════════════════
#  Format A — KM-wise Job Card
# ═════════════════════════════════════════════════════════════════════════════

def _is_jobcard_format(docx_path: Path) -> bool:
    """True only for TrackParametersKmJobCard files."""
    name_lower = docx_path.name.lower()
    return "trackparametersk mjobcard" in name_lower or "kmjobcard" in name_lower or (
        "jobcard" in name_lower and "trackparameters" in name_lower
    )


def _parse_meta_from_filename(docx_path: Path) -> dict:
    """
    Extract metadata from standard filename patterns AND document paragraphs.
    Falls back to folder-name scanning for section / direction hints.
    """
    name = docx_path.name

    # Date in filename: YYYY_MM_DD or YYYY-MM-DD
    m_date = re.search(r"(\d{4})[-_](\d{2})[-_](\d{2})", name)
    run_date = f"{m_date.group(1)}-{m_date.group(2)}-{m_date.group(3)}" if m_date else None

    # TRC no from filename (RT prefix style like RT10201321)
    m_trc_rt = re.search(r"(RT\d+)", name, re.IGNORECASE)
    trc_from_file = m_trc_rt.group(1) if m_trc_rt else None

    # Run no from parent folder name
    m_run = re.search(r"RUN\s*NO\.?\s*[:\.\s]*([A-Z0-9]+)", docx_path.parent.name, re.IGNORECASE)
    run_no_from_folder = m_run.group(1) if m_run else None

    # Section and line from folder name (e.g. "RUN NO. D(IGP-BSL) DN LINE")
    m_sect = re.search(r"\(([A-Z0-9\-]+)\)", docx_path.parent.name, re.IGNORECASE)
    m_line = re.search(r"\b(UP|DN|DOWN)\b", docx_path.parent.name, re.IGNORECASE)
    section_from_folder = m_sect.group(1) if m_sect else None
    line_from_folder    = m_line.group(1).upper() if m_line else None

    return {
        "run_date":  run_date,
        "trc_no":    trc_from_file,
        "run_no":    run_no_from_folder,
        "section":   section_from_folder,
        "line":      line_from_folder,
    }


def _parse_header_meta(row: list[str]) -> dict:
    """
    Parse a metadata table row such as:
      ['TRC NO - 9001', 'DATE: 20.04.2026', 'RUN NO. D', 'ROUTE: A', 'RT-CODE: 10201321', 'FILE NAME: 20042026D']
    or:
      ['RAILWAY: Central', 'DIVISION: BSL', 'SECTION: IGP-BSL', 'START KM: 137', 'LINE: DN']

    Joins all non-empty cells with ' | ' and runs all _META_RE patterns.
    Returns dict with field→clean value (None for missing/dash-only).
    """
    joined = " | ".join(c for c in row if c)
    result = _scan_text_for_meta(joined)
    # Normalise date
    if result.get("run_date"):
        result["run_date"] = _normalise_date(result["run_date"])
    return result


def _extract_section_spd_from_paragraphs(doc) -> str | None:
    """
    Scan document paragraphs for SECTION SPD (KMPH): <value>.
    Returns the first numeric value found, or None.
    """
    pat = re.compile(r"SECTION\s*SPD\s*(?:\(KMPH\))?\s*[:\-]?\s*(\d+)", re.IGNORECASE)
    try:
        for para in doc.paragraphs:
            text = para.text.replace("\xa0", " ").strip()
            if text:
                m = pat.search(text)
                if m:
                    return m.group(1)
    except Exception:
        pass
    return None


def _extract_trc_no_from_paragraphs(doc) -> str | None:
    """
    Scan document paragraphs for the numeric TRC NO (e.g. '- 9001').
    The filename may carry RT10201321 but the readable TRC no is '9001'.
    """
    pat = re.compile(r"(?:CHORD\s*MODE\s*\)\s*-\s*|TRC\s*NO\s*[:\-]\s*)(\d+)", re.IGNORECASE)
    try:
        for para in doc.paragraphs:
            text = para.text.replace("\xa0", " ").strip()
            if text:
                m = pat.search(text)
                if m:
                    return m.group(1)
    except Exception:
        pass
    return None


def convert_jobcard_docx_to_xlsx(docx_path: Path, output_dir: Path) -> Path | None:
    """
    Convert a TrackParametersKmJobCard DOCX (Format A) to XLSX.
    Each KM table → flattened rows, all tables → one sheet.
    """
    log.info(f"    [Format A] KM Job Card: {docx_path.name}")
    try:
        import docx as _docx
        doc = _docx.Document(docx_path)
    except Exception as e:
        log.error(f"    Cannot open DOCX: {e}")
        return None

    body = doc.element.body
    tables_xml = body.findall(f".//{{{_W}}}tbl")
    log.info(f"    Found {len(tables_xml)} KM tables in job card")

    file_meta = _parse_meta_from_filename(docx_path)
    all_rows  = []
    global_meta: dict = {}

    for tbl_idx, tbl in enumerate(tables_xml):
        rows = _table_rows(tbl)
        if len(rows) < 7:
            continue

        if len(rows) > 3:
            meta = _parse_header_meta(rows[3])
            if not global_meta and meta.get("railway"):
                global_meta = meta

        km_meta = _parse_header_meta(rows[3]) if len(rows) > 3 else {}

        for row in rows[7:]:
            if not row or all(c == "" for c in row):
                continue
            row = row + [""] * max(0, 9 - len(row))

            block      = row[0]
            sd         = row[1]
            km         = row[2]
            metre      = row[3]
            uml_param  = row[4]
            uml_value  = row[5]
            cbml_param = row[6]
            cbml_value = row[7]
            track_feat = row[8] if len(row) > 8 else ""

            if not km and not metre and not uml_param and not cbml_param:
                continue

            base = {
                "s_no":          len(all_rows) + 1,
                "block":         block,
                "sd":            sd,
                "km":            km,
                "metre":         metre,
                "track_feature": track_feat,
                "km_range":      km_meta.get("start_km") or km_meta.get("km_range", ""),
                "section":       km_meta.get("section", ""),
                "line":          km_meta.get("line", ""),
                "railway":       km_meta.get("railway", ""),
                "division":      km_meta.get("division", ""),
            }

            if uml_param or uml_value:
                all_rows.append({**base,
                    "parameter":      uml_param,
                    "recorded_value": uml_value,
                    "parameter_type": "UML"})

            if cbml_param or cbml_value:
                all_rows.append({**base,
                    "parameter":      cbml_param,
                    "recorded_value": cbml_value,
                    "parameter_type": "CBML"})

    if not all_rows:
        log.warning(f"    No data rows extracted from {docx_path.name}")
        return None

    df = pd.DataFrame(all_rows)
    log.info(f"    Extracted {len(df)} total parameter rows from job card")

    report_type = "Track Geometry Parameters"
    out_name    = docx_path.stem + "_converted.xlsx"
    out_path    = output_dir / "converted_xlsx" / out_name

    combined_meta = {**global_meta, **file_meta}
    return _write_xlsx(df, out_path, report_type, combined_meta, sheet_name="TrackGeometry")


# ═════════════════════════════════════════════════════════════════════════════
#  Format B — Urgent / Offline / KmReport  (4-table-per-block structure)
# ═════════════════════════════════════════════════════════════════════════════

_URGENT_REPORT_NAMES = (
    "trackparametersurgentreport",
    "trackparametersofflinereport",
    "trackparameterskmreport",
)


def _is_urgent_report_format(docx_path: Path) -> bool:
    """True for Urgent / Offline / KM-Report DOCX files."""
    name_lower = docx_path.name.lower()
    return any(n in name_lower for n in _URGENT_REPORT_NAMES)


def _flatten_loc_header(rows: list[list[str]]) -> list[str]:
    """
    Intelligently flatten the 2-row LOC header used in Urgent/Offline reports.

    Input rows (each is a list of cell strings):
        Row 0: ['PARAMETER', 'THRESHOLD', 'RECORDED VALUE', 'LOC', '']
        Row 1: ['PARAMETER', 'THRESHOLD', 'RECORDED VALUE', 'KM', 'METER']

    Output column names:
        ['parameter', 'threshold', 'recorded_value', 'loc_km', 'loc_meter']

    Algorithm:
        - For each column position, collect unique non-empty labels top-to-bottom.
        - If consecutive labels are same, deduplicate.
        - Join with '_', then snake_case.
        - Special-case: 'loc'+'km' → 'loc_km', 'loc'+'meter' → 'loc_meter'.
    """
    if not rows:
        return []

    n_cols = max(len(r) for r in rows)

    # Pad rows to same width
    padded = [r + [""] * (n_cols - len(r)) for r in rows]

    # Forward-fill parent row (row 0) across columns where row 0 is blank
    # but only for LOC-style parent headers, not for all blank cells
    ff_row0 = list(padded[0])
    last_non_empty = ""
    for i, v in enumerate(ff_row0):
        if v:
            last_non_empty = v
        elif last_non_empty and len(padded) > 1 and padded[1][i]:
            # Fill forward only if row 1 has a child label here
            ff_row0[i] = last_non_empty

    fixed_rows = [ff_row0] + padded[1:]

    result = []
    for col_idx in range(n_cols):
        parts = []
        seen_prev = None
        for row in fixed_rows:
            cell = row[col_idx].strip() if col_idx < len(row) else ""
            cell_lower = cell.lower()
            # Skip empty, duplicate, or placeholder
            if not cell_lower or cell_lower in ("nan", "none", "-"):
                continue
            if cell_lower == seen_prev:
                continue  # deduplicate consecutive same labels
            parts.append(cell_lower)
            seen_prev = cell_lower

        # Join and clean
        raw_name = "_".join(parts).strip("_")
        # Replace spaces and non-alnum
        raw_name = re.sub(r"[^a-z0-9]+", "_", raw_name).strip("_")

        if not raw_name:
            raw_name = f"col_{col_idx}"

        # Canonical mappings for known compound names
        _CANONICAL = {
            "loc_km":           "loc_km",
            "loc_meter":        "loc_meter",
            "loc_metre":        "loc_meter",
            "parameter":        "parameter",
            "threshold":        "threshold",
            "recorded_value":   "recorded_value",
        }
        result.append(_CANONICAL.get(raw_name, raw_name))

    return result


def convert_urgent_report_docx_to_xlsx(docx_path: Path, output_dir: Path) -> Path | None:
    """
    Convert a TrackParametersUrgentReport / OfflineReport / KmReport DOCX (Format B)
    to a clean SQL-ready XLSX.

    The DOCX has a repeating 4-table-per-block structure:
        Block k = (Table 4k, Table 4k+1, Table 4k+2, Table 4k+3)
          4k+0 → metadata row 1: TRC NO, DATE, RUN NO, ROUTE, RT-CODE, FILE NAME
          4k+1 → metadata row 2: RAILWAY, DIVISION, SECTION, START KM, LINE
          4k+2 → 2-row column header: PARAMETER / THRESHOLD / RECORDED VALUE / LOC / (blank)
                                       PARAMETER / THRESHOLD / RECORDED VALUE / KM  / METER
          4k+3 → data rows: ['AGC-M(-)', '-11.00 mm', '-16.56', '137', '47.31']

    Each data row gets:
        parameter, threshold, recorded_value, loc_km, loc_meter  ← from data table
        trc_no, run_date, run_no, route, rt_code, file_name      ← from Table 4k+0
        railway, division, section, start_km, line, section_spd  ← from Table 4k+1 + paragraphs

    Returns path to generated XLSX or None on failure.
    """
    log.info(f"    [Format B] Urgent/Offline/KmReport: {docx_path.name}")
    try:
        import docx as _docx
        doc = _docx.Document(docx_path)
    except Exception as e:
        log.error(f"    Cannot open DOCX: {e}")
        return None

    # Extract section_spd and numeric trc_no from paragraphs (only appears there)
    section_spd_para = _extract_section_spd_from_paragraphs(doc)
    trc_no_para      = _extract_trc_no_from_paragraphs(doc)

    # File-level fallback metadata (from filename + folder)
    file_meta = _parse_meta_from_filename(docx_path)
    # Prefer the readable numeric TRC no from paragraphs over RT-code from filename
    if trc_no_para:
        file_meta["trc_no"] = trc_no_para
    if section_spd_para:
        file_meta["section_spd"] = section_spd_para

    # Read all tables via lxml (avoids InvalidXmlError on missing <w:tblGrid>)
    body      = doc.element.body
    tbls_xml  = body.findall(f".//{{{_W}}}tbl")
    n_tables  = len(tbls_xml)
    log.info(f"    Found {n_tables} tables in document")

    all_rows    : list[dict] = []
    col_names   : list[str]  = []
    block_count              = 0

    # ── Iterate 4-table blocks ────────────────────────────────────────────
    for k in range(0, n_tables, 4):
        if k + 3 >= n_tables:
            # Incomplete block at end — try to process what we have
            break

        t0_rows = _table_rows(tbls_xml[k])     # TRC/DATE/RUN/ROUTE/RTCODE/FILENAME
        t1_rows = _table_rows(tbls_xml[k + 1]) # RAILWAY/DIVISION/SECTION/STARTKM/LINE
        t2_rows = _table_rows(tbls_xml[k + 2]) # 2-row column header
        t3_rows = _table_rows(tbls_xml[k + 3]) # data rows

        # ── Extract block metadata ────────────────────────────────────────
        meta0 = _parse_header_meta(t0_rows[0]) if t0_rows else {}
        meta1 = _parse_header_meta(t1_rows[0]) if t1_rows else {}

        # Merge: file-level fallback fills gaps
        block_meta = {}
        for field in _META_FIELDS:
            # Priority: table-extracted > file/folder-derived
            val = meta0.get(field) or meta1.get(field) or file_meta.get(field)
            block_meta[field] = val

        # Ensure section_spd always comes from paragraphs (it's not in tables)
        if not block_meta.get("section_spd") and section_spd_para:
            block_meta["section_spd"] = section_spd_para

        # ── Detect & flatten column headers (Table 4k+2) ─────────────────
        if not col_names:
            # Use first block's header; validate it looks like a header table
            header_text = " ".join(c for r in t2_rows for c in r).upper()
            if "PARAMETER" in header_text or "THRESHOLD" in header_text:
                col_names = _flatten_loc_header(t2_rows)
                log.info(f"    Column names (block 0): {col_names}")
            else:
                log.warning(f"    Block {k//4}: Table 2 does not look like a header — skipping block")
                continue

        # ── Parse data rows (Table 4k+3) ─────────────────────────────────
        n_data_cols = len(col_names)
        for row in t3_rows:
            if not row or all(c == "" for c in row):
                continue
            # Pad or trim to match column count
            row_padded = (row + [""] * n_data_cols)[:n_data_cols]

            record = dict(zip(col_names, row_padded))
            # Attach block metadata as dedicated columns
            for field in _META_FIELDS:
                record[field] = block_meta.get(field)   # None for missing = SQL NULL

            all_rows.append(record)

        block_count += 1

    log.info(f"    Processed {block_count} blocks, extracted {len(all_rows)} data rows")

    if not all_rows:
        log.warning(f"    No data rows extracted from {docx_path.name}")
        return None

    df = pd.DataFrame(all_rows)

    # Ensure column order: data columns first, then metadata
    data_cols_order = col_names if col_names else []
    meta_cols_order = _META_FIELDS
    final_cols = data_cols_order + [c for c in meta_cols_order if c not in data_cols_order]
    # Only keep columns that actually exist in df
    final_cols = [c for c in final_cols if c in df.columns]
    df = df[final_cols]

    # Determine report type from filename
    report_type = _detect_urgent_report_type(docx_path)
    out_name    = docx_path.stem + "_converted.xlsx"
    out_path    = output_dir / "converted_xlsx" / out_name

    # Use global metadata for the XLSX header block (from first block)
    global_meta = {}
    for field in _META_FIELDS:
        # Use most common non-None value across blocks
        vals = [r[field] for r in all_rows if r.get(field)]
        global_meta[field] = vals[0] if vals else None

    return _write_xlsx(df, out_path, report_type, global_meta, sheet_name="ReportData")


def _detect_urgent_report_type(docx_path: Path) -> str:
    """Classify urgent/offline/km-report DOCX by filename."""
    name = docx_path.name.lower()
    if "urgent" in name:
        return "Track Geometry Urgent Peaks Report"
    if "offline" in name:
        return "Track Geometry Offline Report"
    if "kmreport" in name or "km_report" in name:
        return "Track Geometry KM Report"
    return "Track Geometry Report"


# ═════════════════════════════════════════════════════════════════════════════
#  Format C — Detail Ballast / Fastening / Rail Defect Report
# ═════════════════════════════════════════════════════════════════════════════

def _detect_report_type(docx_path: Path) -> str:
    """Infer report type from filename or folder name."""
    name_lower = (docx_path.name + " " + docx_path.parent.name).lower()
    if "ballast" in name_lower:
        return "Detail Ballast Report"
    if "fastening" in name_lower or "fittings" in name_lower:
        return "Detail Fastening Report"
    if "rail" in name_lower and "defect" in name_lower:
        return "Rail Defect Detail Report"
    return "Detail Inspection Report"


def convert_detail_report_docx_to_xlsx(docx_path: Path, output_dir: Path) -> Path | None:
    """
    Convert a Detail Ballast/Fastening/Rail DOCX (Format C) to XLSX.
    Paragraphs are parsed for LOC + defect pairs.
    """
    log.info(f"    [Format C] Detail Report: {docx_path.name}")
    try:
        import docx as _docx
        doc = _docx.Document(docx_path)
    except Exception as e:
        log.error(f"    Cannot open DOCX: {e}")
        return None

    # ── Extract Railway/Division/Section from first table (if present) ──────
    body      = doc.element.body
    tbls_xml  = body.findall(f".//{{{_W}}}tbl")
    header_meta: dict = {}
    if tbls_xml:
        first_rows = _table_rows(tbls_xml[0])
        if first_rows:
            header_meta = _parse_header_meta(first_rows[0])

    # ── Extract TRC/date from filename ───────────────────────────────────────
    file_meta   = _parse_meta_from_filename(docx_path)
    report_type = _detect_report_type(docx_path)

    section = header_meta.get("section") or file_meta.get("section", "")
    line    = header_meta.get("line")    or file_meta.get("line", "")

    # ── Parse paragraphs ─────────────────────────────────────────────────────
    records      = []
    current_loc  = None
    current_type = ""

    def _clean(t: str) -> str:
        return t.replace("\x00", "").replace("\xa0", " ").strip()

    for p in doc.paragraphs:
        text = _clean(p.text)
        if not text:
            continue

        if text in ("Detail Ballast Report", "Detail Fastening Report",
                    "Rail Defect Detail Report", "Detail Inspection Report"):
            current_type = text
            continue

        if text == "Details:":
            continue

        if text.startswith("LOC:"):
            m_line_  = re.search(r"Line\s*[-:=]\s*([A-Za-z0-9]+)", text, re.IGNORECASE)
            m_km     = re.search(r"KM\s*[^\d]*(\d+)", text, re.IGNORECASE)
            m_meter  = re.search(r"Meter\s*[^\d]*([\d.]+)", text, re.IGNORECASE)
            m_rail   = re.search(r"Rail\s*[-:=]\s*([A-Za-z]+)", text, re.IGNORECASE)
            current_loc = {
                "line":       m_line_.group(1)  if m_line_  else line,
                "loc_km":     m_km.group(1)     if m_km     else None,
                "loc_meter":  m_meter.group(1)  if m_meter  else None,
                "rail_side":  m_rail.group(1)   if m_rail   else None,
            }
            continue

        if current_loc is not None:
            records.append({
                "s_no":            len(records) + 1,
                "report_type":     current_type or report_type,
                "line":            current_loc["line"],
                "loc_km":          current_loc["loc_km"],
                "loc_meter":       current_loc["loc_meter"],
                "rail_side":       current_loc["rail_side"],
                "defect_type":     text,
                "section":         section,
                "railway":         _clean_meta_value(header_meta.get("railway")),
                "division":        _clean_meta_value(header_meta.get("division")),
                "trc_no":          file_meta.get("trc_no"),
                "run_date":        file_meta.get("run_date"),
                "source_file":     docx_path.name,
            })
            current_loc = None

    if not records:
        log.warning(f"    No LOC records found in {docx_path.name}")
        return None

    df = pd.DataFrame(records)
    log.info(f"    Extracted {len(df)} defect records")

    out_name = docx_path.stem + "_converted.xlsx"
    out_path = output_dir / "converted_xlsx" / out_name

    combined_meta = {**header_meta, **file_meta, "section": section, "line": line}
    return _write_xlsx(df, out_path, report_type, combined_meta, sheet_name="DefectRecords")


# ═════════════════════════════════════════════════════════════════════════════
#  Public dispatcher
# ═════════════════════════════════════════════════════════════════════════════

def convert_docx_to_xlsx(docx_path: Path, output_dir: Path) -> Path | None:
    """
    Detect DOCX format and run the appropriate converter.

    Routing priority:
        1. KM Job Card format    → convert_jobcard_docx_to_xlsx()
        2. Urgent/Offline/KM    → convert_urgent_report_docx_to_xlsx()
        3. Everything else       → convert_detail_report_docx_to_xlsx()

    Args:
        docx_path  : Path to the source .docx file
        output_dir : Root folder; converted XLSX written to output_dir/converted_xlsx/

    Returns:
        Path to the generated XLSX, or None on failure.
    """
    if not docx_path.exists():
        log.error(f"DOCX not found: {docx_path}")
        return None

    try:
        if _is_jobcard_format(docx_path):
            return convert_jobcard_docx_to_xlsx(docx_path, output_dir)
        elif _is_urgent_report_format(docx_path):
            return convert_urgent_report_docx_to_xlsx(docx_path, output_dir)
        else:
            return convert_detail_report_docx_to_xlsx(docx_path, output_dir)
    except Exception as ex:
        log.error(f"    Conversion failed for {docx_path.name}: {ex}", exc_info=True)
        return None
