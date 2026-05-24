"""
Diagnostic script: inspect the raw DOCX structure for TrackParametersUrgentReport
to understand what metadata and table content looks like before conversion.
"""
import re
from pathlib import Path

try:
    import docx as _docx
except ImportError:
    print("python-docx not installed")
    exit()

# Pick the urgent report
docx_path = Path("data/April-26 TRC-9001/RUN NO. D(IGP-BSL) DN LINE/TrackParametersUrgentReport_2026_04_20_20_31_24_361.docx")

print(f"Opening: {docx_path.name}")
doc = _docx.Document(docx_path)

print(f"\n=== PARAGRAPHS (first 40) ===")
for i, p in enumerate(doc.paragraphs[:40]):
    txt = p.text.replace('\xa0', ' ').strip()
    if txt:
        print(f"  [{i:3d}] {repr(txt)}")

print(f"\n=== TABLES ({len(doc.tables)} total) ===")
for t_idx, tbl in enumerate(doc.tables[:3]):
    print(f"\n--- Table {t_idx} ---")
    for r_idx, row in enumerate(tbl.rows[:10]):
        cells = [c.text.replace('\xa0', ' ').strip() for c in row.cells]
        print(f"  Row {r_idx}: {cells}")
    if len(tbl.rows) > 10:
        print(f"  ... ({len(tbl.rows)} total rows)")

# Also check the XML body for raw table structure
from lxml import etree
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
body = doc.element.body
tbls_xml = body.findall(f".//{{{_W}}}tbl")
print(f"\n=== XML TABLES ({len(tbls_xml)} via lxml) ===")
for i, tbl in enumerate(tbls_xml[:2]):
    rows = tbl.findall(f"{{{_W}}}tr")
    print(f"  Table {i}: {len(rows)} rows")
    for r_idx, tr in enumerate(rows[:8]):
        cells = []
        for tc in tr.findall(f"{{{_W}}}tc"):
            texts = tc.findall(f".//{{{_W}}}t")
            cell_text = "".join(t.text or "" for t in texts).replace("\xa0", " ").strip()
            cells.append(cell_text)
        print(f"    Row {r_idx}: {cells}")
