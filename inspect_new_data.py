"""
Inspect new data files: CSVs, DOCX tables, and MDB structure.
"""
import os, sys
import pandas as pd
from pathlib import Path

BASE = Path(r"d:\ENGINEER\IndianRailwaysProject - Second_Data\data\TRC Runs BSL Div DVD Data")

print("=" * 70)
print("SCANNING ALL DATA FILES IN data/ FOLDER")
print("=" * 70)

# Walk all files
all_files = []
for root, dirs, files in os.walk(BASE):
    for f in files:
        fpath = Path(root) / f
        all_files.append(fpath)
        rel = fpath.relative_to(BASE)
        size_mb = fpath.stat().st_size / (1024*1024)
        print(f"  [{fpath.suffix.upper():6}] {size_mb:7.1f} MB  |  {rel}")

print(f"\nTotal files: {len(all_files)}")

# Separate by type
csvs  = [f for f in all_files if f.suffix.lower() == '.csv']
docxs = [f for f in all_files if f.suffix.lower() == '.docx']
mdbs  = [f for f in all_files if f.suffix.lower() == '.mdb']

print(f"\nCSV  : {len(csvs)}")
print(f"DOCX : {len(docxs)}")
print(f"MDB  : {len(mdbs)}")

# ── Inspect CSVs ──────────────────────────────────────────────────────────────
print("\n" + "=" * 70)
print("CSV FILE SCHEMAS")
print("=" * 70)

for csv_path in csvs:
    print(f"\n>>> {csv_path.name}")
    try:
        # Read just header + few rows, detect encoding
        df = pd.read_csv(csv_path, nrows=5, encoding='utf-8', low_memory=False)
    except UnicodeDecodeError:
        try:
            df = pd.read_csv(csv_path, nrows=5, encoding='latin-1', low_memory=False)
        except Exception as ex:
            print(f"    ERROR reading: {ex}")
            continue
    except Exception as ex:
        print(f"    ERROR: {ex}")
        continue

    print(f"    Shape (sample): {df.shape}")
    print(f"    Columns ({len(df.columns)}): {list(df.columns)}")
    print(f"    Dtypes: {df.dtypes.to_dict()}")
    print(f"    Sample row 0: {df.iloc[0].to_dict() if len(df) > 0 else 'EMPTY'}")

# ── Inspect DOCX ──────────────────────────────────────────────────────────────
print("\n" + "=" * 70)
print("DOCX FILE TABLE INSPECTION (first 2 tables, 3 rows each)")
print("=" * 70)

try:
    from docx import Document
    DOCX_OK = True
except ImportError:
    print("  python-docx not installed. Run: pip install python-docx")
    DOCX_OK = False

if DOCX_OK:
    for docx_path in docxs:
        print(f"\n>>> {docx_path.name}")
        try:
            doc = Document(str(docx_path))
            print(f"    Paragraphs: {len(doc.paragraphs)}")
            print(f"    Tables    : {len(doc.tables)}")
            # Show first few paragraphs (metadata)
            for i, para in enumerate(doc.paragraphs[:8]):
                txt = para.text.strip()
                if txt:
                    print(f"    PARA[{i}]: {txt[:120]}")
            # Show first 2 tables
            for ti, tbl in enumerate(doc.tables[:2]):
                print(f"    TABLE[{ti}]: {len(tbl.rows)} rows x {len(tbl.columns)} cols")
                for ri, row in enumerate(tbl.rows[:3]):
                    cells = [c.text.strip() for c in row.cells]
                    print(f"      ROW[{ri}]: {cells}")
        except Exception as ex:
            print(f"    ERROR: {ex}")

print("\n=== DONE ===")
